#!/usr/bin/env python3
"""
Gera documento DOCX COMPLETO: 20 Projetos de Pesquisa em Reproducao Equina
PARTE A: 10 projetos de Diluentes para Semen Equino
PARTE B: 10 projetos de Capacitacao Espermatica para FIV Equina
Rastreio de literatura + justificativa de diferencial
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

def add_table_row(table, cells_text, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        if bold:
            run.bold = True
    return row

def set_table_style(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)

def make_header_row(table):
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)

def write_project(doc, proj, chapter_num):
    """Write a single project section."""
    doc.add_page_break()
    doc.add_heading(f'{chapter_num}. Projeto {proj["id"]} — {proj["title"]}', level=1)

    doc.add_heading('Rastreio de Literatura', level=2)
    for ref_author, ref_journal, ref_desc in proj['literature']:
        p = doc.add_paragraph()
        run = p.add_run(f'{ref_author} ')
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(f'({ref_journal})')
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        p2 = doc.add_paragraph(ref_desc)
        p2.paragraph_format.left_indent = Cm(1)
        p2.paragraph_format.space_after = Pt(8)

    doc.add_heading('Gap Identificado na Literatura', level=2)
    doc.add_paragraph(proj['gap'])

    doc.add_heading('Diferencial e Justificativa de Originalidade', level=2)
    doc.add_paragraph(proj['differential'])

    doc.add_heading('Desenho Experimental', level=2)
    for line in proj['design'].split('\n'):
        stripped = line.strip()
        if stripped:
            if stripped.startswith('-'):
                doc.add_paragraph(stripped[1:].strip(), style='List Bullet')
            else:
                doc.add_paragraph(stripped)
        else:
            doc.add_paragraph()

    doc.add_heading('Equipamentos e Reagentes Necessarios', level=2)
    doc.add_paragraph(proj['equipment'])

    doc.add_heading('Estimativa de Custo', level=2)
    doc.add_paragraph(proj['cost'])

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROJETOS DE PESQUISA EM\nREPRODUCAO EQUINA')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('20 Propostas Experimentais para a UNESP Botucatu')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2E, 0x5E, 0x8E)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Parte A: Diluentes para Semen Equino (10 projetos)\nParte B: Capacitacao Espermatica para FIV Equina (10 projetos)')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Rastreio de Literatura e Justificativa de Diferencial')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.italic = True

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Elaborado com auxilio do Equine Reproduction Research Squad\n'
                'Especialistas consultados: Papa, Pena, Dell\'Aqua Jr, Alvarenga, Sieme, Gibb,\n'
                'Leemans, Hinrichs, Felix, Maitan, Stout, Squires, Morrell')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Abril 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Sumario', level=1)
toc = [
    '1. Introducao e Contexto Geral',
    '2. Justificativa — Vantagens Competitivas da UNESP',
    '',
    'PARTE A — DILUENTES PARA SEMEN EQUINO',
    '3. Estado da Arte: Diluentes e Criopreservacao (2024-2026)',
    '4. Projeto A1 — Piruvato no BotuSemen para Refrigeracao 48-72h',
    '5. Projeto A2 — Trehalose 50 mM no BotuCrio com Taxa de Prenhez',
    '6. Projeto A3 — BotuCrio + DMF Exclusivo em Bad Freezers com IA',
    '7. Projeto A4 — Plasma Seminal Dose-Resposta x Tipo de Garanhao',
    '8. Projeto A5 — BHT Convencional vs BHT Nanoparticulado no BotuSemen',
    '9. Projeto A6 — SpermFilter vs SLC para Selecao Pre-Congelacao',
    '10. Projeto A7 — Re-suspensao Pos-Refrigeracao Aplicada a Congelacao',
    '11. Projeto A8 — Fracionamento do Ejaculado x Qualidade Pos-Congelacao',
    '12. Projeto A9 — BotuCrio Liofilizado: Prova de Conceito',
    '13. Projeto A10 — Piruvato + L-Carnitina no BotuSemen para 72h',
    '14. Sintese Parte A — Priorizacao dos Projetos de Diluentes',
    '',
    'PARTE B — CAPACITACAO ESPERMATICA PARA FIV EQUINA',
    '15. Estado da Arte: FIV Equina e Capacitacao (2022-2026)',
    '16. Projeto B1 — Reproducao do Protocolo PHE com BotuCrio',
    '17. Projeto B2 — Cinetica de Capacitacao x Tipo de Crioprotetor',
    '18. Projeto B3 — SpermFilter vs Swim-up vs SSD Pre-FIV',
    '19. Projeto B4 — Plasma Seminal Residual e Capacitacao In Vitro',
    '20. Projeto B5 — Piruvato no Meio de Capacitacao PHE',
    '21. Projeto B6 — Termotaxia como Selecao Pre-FIV',
    '22. Projeto B7 — pH Alcalino (7,9) no Meio de Capacitacao',
    '23. Projeto B8 — Bad Freezers na FIV: Criocapacitacao como Vantagem?',
    '24. Projeto B9 — Antioxidante no Meio de Capacitacao',
    '25. Projeto B10 — Metabolitos Oviductais no Meio PHE',
    '26. Sintese Parte B — Priorizacao dos Projetos de Capacitacao/FIV',
    '',
    '27. Proposta de Integracao: Duas Teses de Doutorado',
    '28. Referencias Bibliograficas',
]
for item in toc:
    if item == '':
        doc.add_paragraph()
    elif item.startswith('PARTE'):
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    else:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. INTRODUCAO
# ============================================================
doc.add_heading('1. Introducao e Contexto Geral', level=1)

doc.add_paragraph(
    'Este documento apresenta 20 propostas experimentais em reproducao equina, organizadas em '
    'duas linhas de pesquisa complementares: (A) diluentes para semen equino — refrigeracao e '
    'criopreservacao; e (B) capacitacao espermatica para fertilizacao in vitro (FIV) convencional. '
    'Todas as propostas foram desenhadas para serem viaveis com a infraestrutura da FMVZ/UNESP '
    'Botucatu e representam gaps genuinos na literatura — experimentos que nao foram realizados '
    'pela equipe UNESP nem por outros grupos.'
)
doc.add_paragraph(
    'Cada projeto inclui: (1) rastreio completo do que existe na literatura; (2) identificacao '
    'precisa do gap; (3) justificativa do diferencial de originalidade; (4) desenho experimental '
    'detalhado; (5) lista de equipamentos e reagentes; e (6) estimativa de custo.'
)
doc.add_paragraph(
    'As duas linhas convergem naturalmente: a Parte A otimiza a criopreservacao do semen, e a '
    'Parte B usa esse semen otimizado para FIV. Juntas, formam um programa de pesquisa coerente '
    'que posiciona a UNESP na fronteira da reproducao equina assistida.'
)

doc.add_page_break()

# ============================================================
# 2. JUSTIFICATIVA ��� UNESP
# ============================================================
doc.add_heading('2. Justificativa — Vantagens Competitivas da UNESP', level=1)

advantages = [
    ('BotuCrio e BotuSemen (Botupharma)',
     'Diluentes proprios com ampla distribuicao comercial mundial. Possibilidade de testar '
     'formulacoes proprietarias e posicionar produtos para novos mercados (FIV).'),
    ('SpermFilter',
     'Tecnologia propria de separacao espermatica com 89,4% de recuperacao. Potencial para '
     'mercado de FIV se demonstrada eficacia como metodo de selecao pre-FIV.'),
    ('Rebanho experimental classificado',
     'Garanhoes classificados como bons, intermediarios e maus congeladores. Eguas disponIveis '
     'para validacao in vivo (taxa de prenhez). Infraestrutura que poucos centros no mundo possuem.'),
    ('CASA + Citometria de fluxo',
     'Avaliacao multiparametrica: motilidade (CASA), integridade de membrana (CFDA/PI), potencial '
     'mitocondrial (JC-1), producao de ROS (DHE), capacitacao (CTC).'),
    ('Acesso a oocitos de abatedouro',
     'Ovarios de eguas abatidas para obtencao de oocitos para testes in vitro sem necessidade '
     'de OPU em eguas vivas — reduz custo e complexidade.'),
    ('Expertise em criopreservacao',
     'Grupo Papa/Alvarenga/Dell\'Aqua Jr com decadas de publicacoes e reconhecimento internacional. '
     'Entendimento profundo de criocapacitacao e variabilidade entre garanhoes.'),
    ('Nenhuma publicacao em FIV convencional',
     'A UNESP ainda nao publicou dados de FIV equina convencional — oportunidade de ser o '
     'primeiro grupo brasileiro nessa linha, apos o breakthrough de Hinrichs (2022).'),
    ('Potencial de colaboracao interdepartamental',
     'Instituto de Quimica (sintese de nanoparticulas), Instituto de Biociencias (liofilizador), '
     'FMVZ (reproducao) — campus integrado para projetos multidisciplinares.'),
]
for title, desc in advantages:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ============================================================
# PARTE A — DILUENTES
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PARTE A')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Diluentes para Semen Equino')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2E, 0x5E, 0x8E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Refrigeracao e Criopreservacao — 10 Projetos')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ============================================================
# 3. ESTADO DA ARTE — DILUENTES
# ============================================================
doc.add_heading('3. Estado da Arte: Diluentes e Criopreservacao (2024-2026)', level=1)

doc.add_heading('3.1 O que a UNESP ja publicou (para nao repetir)', level=2)

table = doc.add_table(rows=1, cols=2, style='Table Grid')
table.rows[0].cells[0].text = 'Tema ja publicado'
table.rows[0].cells[1].text = 'Referencia'
make_header_row(table)

published = [
    ('BotuCrio (MF + glicerol + gema 20%)', 'Papa, Alvarenga — padrao comercial'),
    ('DMF vs glicerol em bad freezers', 'Papa — amidas > glicerol'),
    ('SpermFilter vs centrifugacao', 'Papa — SpermFilter 89,4% recuperacao'),
    ('Caixa de isopor = maquinas automatizadas', 'Maziero/Papa 2010'),
    ('Colesterol/ciclodextrina (in vitro)', 'Grupo Botucatu — sem dados in vivo'),
    ('Post-cooling re-suspension (bad coolers)', 'Segabinazzi, Dell\'Aqua, Papa 2024'),
    ('Ejaculacao farmacologica', 'Papa — detomidina + ocitocina'),
    ('Vesiculite seminal — tratamento endoscopico', 'Papa'),
    ('Semen epididimario — criopreservacao', 'Papa/Alvarenga'),
    ('LDL purificada vs gema de ovo no BotuCrio', 'Publicado 2020 — inconclusivo'),
    ('L-carnitina + acetil-L-carnitina a 5 graus C', 'Publicado com BotuSemen'),
]
for tema, ref in published:
    add_table_row(table, [tema, ref])
set_table_style(table)

doc.add_paragraph()

doc.add_heading('3.2 Estudos recentes de outros grupos (2024-2026)', level=2)

recent = [
    ('Nanoparticulas de BHT (2025)',
     'BHT-NPs a 0,5-1,0 mM melhoraram motilidade, viabilidade e membrana em semen refrigerado '
     'a 4 graus C por 72h. Tripla acao: antioxidante + antimicrobiana + anti-apoptotica. '
     'Springer, Vet Res Commun, 2025.'),
    ('Nanoparticulas de selenio em INRA96 (2024)',
     'SeNPs em INRA96 para garanhoes Turcomanos: melhora de motilidade, membrana, reducao de '
     'peroxidacao lipidica. Theriogenology, 2024.'),
    ('Beyond — diluente quimicamente definido (2025)',
     'Diluente sem produtos animais. 61% de taxa de embrioes com semen armazenado por 5,5 dias '
     'a 5 graus C. Promete ate 14 dias. Reproduction in Domestic Animals, 2025.'),
    ('Glicerol vs amidas x curvas de resfriamento (2024)',
     'Fatorial crioprotetor x curva: interacao significativa. MF preserva melhor funcao '
     'mitocondrial. Theriogenology, 2024.'),
    ('Trehalose 50 mM em diluente de lactose (2023)',
     '50 mM otimo — melhora motilidade, acrossoma, membrana, DNA. 150 mM deleteria. '
     'Tendencia de melhor prenhez (n insuficiente). J Equine Vet Sci, 2023.'),
    ('Diluentes sem antibioticos (2025)',
     'Filtracao e SLC mantiveram qualidade comparavel a protocolo com antibioticos. '
     'Peptideos antimicrobianos como alternativa. Animals, 2025.'),
    ('Diluentes liofilizados (2024)',
     'Qualidade pos-descongelacao equivalente a diluentes convencionais. Elimina cadeia de frio '
     'para o diluente. Theriogenology, 2024.'),
    ('Extratos naturais — RESULTADO NEGATIVO (2025)',
     'Matcha, spirulina, rabano e quercetina NAO melhoraram congelabilidade. '
     'Variabilidade individual mascarou efeitos. Antioxidantes genericos fracassam. '
     'Antioxidants, 2025.'),
]
for title, desc in recent:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(desc)

doc.add_page_break()

# ============================================================
# PROJETOS A1-A10 (DILUENTES)
# ============================================================
diluent_projects = [
    {
        'id': 'A1',
        'title': 'Piruvato no BotuSemen para Refrigeracao 48-72h',
        'literature': [
            ('Pena et al. (2021)', 'Biol Reprod, PMID 33864078',
             'Low glucose (1 mM) + high pyruvate (10 mM): motilidade 76,2% vs 61,7% (INRA 96). '
             'Reducao de 2-oxoaldeidos (metilglioxal). Mitocondrias ativas: 51,1% vs 24,1% apos 48h a 18 graus C.'),
            ('Pena et al. (2023)', 'Theriogenology, PMID 38029686',
             'Piruvato melhora funcao espermatica em meios com alta glicose. Mecanismo: conversao '
             'piruvato->lactato regenera NAD+ para glicolise. Piruvato alimenta Krebs diretamente.'),
            ('Gibb et al. (2016)', 'BioMed Research International',
             'Lactato e piruvato sao as fontes de energia mais importantes para motilidade espermatica '
             'equina, com efeitos dose-dependentes na funcao mitocondrial.'),
            ('Aging of stallion spermatozoa (2023)', 'Andrology',
             'Meio com 67 mM glicose + 10 mM piruvato atrasa envelhecimento espermatico a 22 graus C.'),
        ],
        'gap': 'O Pena demonstrou que piruvato melhora funcao mitocondrial dramaticamente em meios '
               'de armazenamento genericos. A UNESP NUNCA testou piruvato como suplemento no BotuSemen. '
               'Nenhum estudo publicado combinou piruvato com BotuSemen especificamente, e nenhum testou '
               'a dose-resposta de piruvato em refrigeracao a 5 graus C por 48-72h.',
        'differential': 'Primeiro estudo a testar piruvato no BotuSemen — o diluente de refrigeracao '
                       'mais utilizado comercialmente no Brasil. Se piruvato melhorar a longevidade '
                       'espermatica no BotuSemen, e uma otimizacao simples e barata com impacto '
                       'comercial direto para a Botupharma.',
        'design': (
            'Split-sample, 4 grupos:\n'
            '- G1: BotuSemen padrao (controle)\n'
            '- G2: BotuSemen + piruvato 2 mM\n'
            '- G3: BotuSemen + piruvato 5 mM\n'
            '- G4: BotuSemen + piruvato 10 mM\n\n'
            'Avaliacao: 0h, 24h, 48h, 72h a 5 graus C.\n\n'
            'Parametros:\n'
            '- Motilidade total e progressiva (CASA)\n'
            '- Potencial mitocondrial (JC-1) — endpoint primario\n'
            '- Integridade de membrana (CFDA/PI)\n'
            '- Producao de ROS (DHE)\n'
            '- Morfologia\n\n'
            'n = 10 garanhoes (5 bons coolers + 5 bad coolers).'
        ),
        'equipment': 'BotuSemen (disponivel), piruvato de sodio (~R$50-80/frasco), CASA, citometria (JC-1, DHE, CFDA/PI).',
        'cost': 'Muito baixo — reagente barato + consumiveis de rotina.',
    },
    {
        'id': 'A2',
        'title': 'Trehalose 50 mM no BotuCrio — Validacao com Taxa de Prenhez',
        'literature': [
            ('Trehalose em diluente de lactose (2023)', 'J Equine Vet Sci',
             '50 mM de trehalose: melhora motilidade progressiva, integridade acrossomal, membrana, '
             'potencial mitocondrial e DNA pos-descongelacao. 150 mM: efeito osmotico deletero.'),
            ('Trehalose review — pequenos ruminantes (2024)', 'Frontiers Vet Sci',
             'Trehalose protege contra estresse osmotico, desidrata espermatozoides antes da congelacao, '
             'protege contra dano oxidativo. Concentracao e especie-especifica.'),
            ('Papa — hierarquia de evidencia', 'Diversos UNESP',
             'Taxa de prenhez in vivo e o unico gold standard. Motilidade e proxy, nao preditor absoluto. '
             'Papa exige validacao dupla (lab + campo) antes de aceitar novas tecnologias.'),
        ],
        'gap': 'Trehalose 50 mM no diluente de congelacao ja mostrou melhora in vitro (outro grupo, '
               'diluente a base de lactose). NINGUEM validou in vivo com taxa de prenhez. '
               'NINGUEM testou trehalose no BotuCrio especificamente.',
        'differential': 'Primeiro estudo a validar trehalose com o endpoint que realmente importa: '
                       'taxa de prenhez por ciclo. A UNESP e o unico grupo com acesso a BotuCrio '
                       '+ rebanho experimental de eguas para IA na mesma estrutura.',
        'design': (
            '- G1: BotuCrio padrao (controle)\n'
            '- G2: BotuCrio + trehalose 50 mM\n\n'
            '40 ciclos/grupo (poder para detectar diferenca de 15%).\n'
            'Garanhoes pareados entre grupos (mesmos garanhoes nos 2 grupos).\n\n'
            'Endpoint primario: taxa de prenhez/ciclo (US dia 15).\n'
            'Endpoints secundarios: CASA T0/T30, membrana (CFDA/PI), JC-1.\n\n'
            'n = 8-10 garanhoes, 80 ciclos totais.'
        ),
        'equipment': 'BotuCrio (disponivel), trehalose (reagente barato), eguas experimentais (disponivel), ultrassom.',
        'cost': 'Moderado — principalmente custo de manejo reprodutivo das eguas.',
    },
    {
        'id': 'A3',
        'title': 'BotuCrio + DMF Exclusivo em Bad Freezers — com Inseminacao Artificial',
        'literature': [
            ('Papa et al. — amidas vs glicerol', 'Diversos UNESP',
             'DMF superior ao glicerol em 40/55 garanhoes. Mangalarga Marchador: 40% prenhez com DMF '
             'vs 0/15 com glicerol. Amidas tem menor peso molecular e melhor permeabilidade.'),
            ('Unraveling glycerol and amides (2024)', 'Theriogenology',
             'Fatorial GLY vs MF vs DMF vs DMA x curvas. MF preserva melhor funcao mitocondrial. '
             'Interacao crioprotetor x curva e significativa.'),
            ('Alvarenga — classificacao de garanhoes', 'UNESP',
             'Bad freezer: <20% motilidade progressiva pos-descongelacao. Protocolo padrao: '
             'MF + glicerol (BotuCrio). Alternativas para bad freezers: trocar crioprotetor.'),
        ],
        'gap': 'A UNESP ja mostrou que amidas > glicerol. Mas NUNCA comparou BotuCrio padrao '
               '(MF + glicerol) vs uma versao com DMF exclusivo (sem glicerol) EXCLUSIVAMENTE '
               'em bad freezers E com dados de prenhez.',
        'differential': 'Primeiro estudo exclusivamente em bad freezers comparando combinacao padrao '
                       'vs amida pura no BotuCrio com taxa de prenhez como endpoint primario.',
        'design': (
            'Somente bad freezers (n = 8-10 garanhoes, <20% motilidade pos-descongelacao padrao).\n\n'
            '- G1: BotuCrio padrao (MF + glicerol)\n'
            '- G2: BotuCrio modificado (DMF 5% exclusivo, sem glicerol)\n'
            '- G3: BotuCrio modificado (DMF 3% + MF 2%)\n\n'
            '15-20 ciclos/grupo.\n'
            'Endpoint primario: taxa de prenhez.\n'
            'Endpoint secundario: CASA T0/T30, membrana, DNA.'
        ),
        'equipment': 'BotuCrio base + DMF + MF (reagentes disponIveis), eguas para IA.',
        'cost': 'Moderado.',
    },
    {
        'id': 'A4',
        'title': 'Plasma Seminal Dose-Resposta x Classificacao do Garanhao',
        'literature': [
            ('Sieme et al.', 'TiHo Hannover, diversos',
             'Plasma seminal: 5-10% otimo para CRIOPRESERVACAO. >50% toxico. Componentes protetores: '
             'HSPs, clusterin, lactoferrin, enzimas antioxidantes. Dose-resposta: 0% (desprotegido), '
             '5-10% (otimo), 20-30% (variavel), >50% (toxico).'),
            ('Sieme — classificacao individual', 'Diversos',
             '30% bons (>40% motilidade pos-descongelacao), 40% intermediarios, 30% maus (<20%). '
             'Composicao do plasma difere entre bons e maus congeladores.'),
            ('Papa — SpermFilter', 'UNESP',
             'SpermFilter remove ~95% do plasma. Centrifugacao 600xg remove ~85-90%. '
             'Nivel residual depende do metodo.'),
        ],
        'gap': 'NINGUEM na UNESP comparou sistematicamente 0%, 5%, 10% e 20% de plasma residual '
               'CRUZANDO com a classificacao do garanhao (bom/intermediario/mau congelador) no BotuCrio. '
               'O Sieme mostrou o conceito, mas com diluentes europeus, nao com BotuCrio.',
        'differential': 'Primeiro estudo fatorial plasma x classificacao no contexto BotuCrio. '
                       'Hipotese: bad freezers podem ter plasma deletero — dose ideal pode ser ZERO. '
                       'Bons congeladores podem tolerar mais plasma.',
        'design': (
            '20 garanhoes pre-classificados (bom/intermediario/mau).\n'
            'Split-sample: 4 doses de plasma (0%, 5%, 10%, 20%).\n'
            'Controle via centrifugacao + reconstituicao com plasma autologico.\n\n'
            'Diluicao final em BotuCrio -> congelacao padrao.\n'
            'Avaliacao pos-descongelacao: CASA T0/T30, membrana, JC-1.\n'
            'Analise: interacao plasma x classificacao (fatorial 4x3).\n\n'
            'n = 20 garanhoes x 4 doses = 80 amostras.'
        ),
        'equipment': 'SpermFilter ou centrifugacao (disponivel), BotuCrio, CASA, citometria.',
        'cost': 'Baixo — so consumiveis.',
    },
    {
        'id': 'A5',
        'title': 'BHT Convencional vs BHT Nanoparticulado no BotuSemen',
        'literature': [
            ('BHT-NPs (2025)', 'Springer, Vet Res Commun',
             'BHT-NPs a 0,5-1,0 mM: melhora significativa de motilidade, viabilidade, membrana '
             'apos 72h a 4 graus C. MEV: reducao de dano de cabeca, peca intermediaria e caudas '
             'enroladas. Tripla acao: antioxidante + antimicrobiana + anti-apoptotica.'),
            ('Nanoparticulas de selenio em INRA96 (2024)', 'Theriogenology',
             'SeNPs melhoram motilidade e reduzem peroxidacao em garanhoes Turcomanos a 5 graus C.'),
            ('Nanotecnologia em semen — editorial (2026)', 'Frontiers Vet Sci',
             'Nanoparticulas como fronteira transformadora para qualidade seminal em pecuaria.'),
        ],
        'gap': 'O estudo de BHT-NPs (2025) usou diluente generico, NAO BotuSemen. Ninguem comparou '
               'BHT convencional vs BHT nanoparticulado no mesmo experimento equino. '
               'O Papa ja pesquisa BHT convencional — mas sem forma nanoparticulada.',
        'differential': 'Primeiro head-to-head de BHT convencional vs BHT-nano no BotuSemen. '
                       'Possibilidade de colaboracao com Instituto de Quimica da UNESP para sintese '
                       'das nanoparticulas — projeto interdisciplinar.',
        'design': (
            'Split-sample, 4 grupos:\n'
            '- G1: BotuSemen (controle)\n'
            '- G2: BotuSemen + BHT convencional 1 mM\n'
            '- G3: BotuSemen + BHT-NPs 0,5 mM\n'
            '- G4: BotuSemen + BHT-NPs 1,0 mM\n\n'
            'Avaliacao: 0h, 24h, 48h, 72h a 5 graus C.\n'
            'Parametros: CASA, membrana, ROS (DHE), morfologia (MEV se disponivel).\n\n'
            'n = 10 garanhoes.'
        ),
        'equipment': 'BotuSemen, BHT convencional (barato), BHT-NPs (sintese via parceria com Quimica), CASA, citometria.',
        'cost': 'Baixo-moderado (depende da sintese das NPs).',
    },
    {
        'id': 'A6',
        'title': 'SpermFilter vs SLC (Single Layer Centrifugation) Pre-Congelacao',
        'literature': [
            ('Papa — SpermFilter', 'Diversos UNESP',
             'Recuperacao 89,4% (vs 80,9% centrifugacao). Sem dano centrifugo. Reutilizavel.'),
            ('Morrell — SLC', 'SLU, Suecia, diversos',
             'Androcoll-E para selecao por motilidade e morfologia. Usado em bovinos para FIV. '
             'Em equinos: melhora motilidade e morfologia pos-selecao. Testado para antibiotico-free.'),
            ('Antibiotic-free approaches (2025)', 'Animals',
             'SLC manteve qualidade comparavel a protocolo com antibioticos em semen equino '
             'criopreservado. Alternativa sustentavel.'),
        ],
        'gap': 'NINGUEM comparou SpermFilter vs SLC no contexto da criopreservacao equina com '
               'BotuCrio. SpermFilter e SLC operam por mecanismos diferentes (filtracao vs selecao '
               'por gradiente) — podem ter resultados distintos em bad freezers.',
        'differential': 'Comparacao direta entre tecnologia brasileira (SpermFilter/UNESP) e '
                       'europeia (SLC/Morrell). Posiciona o SpermFilter no mercado global de '
                       'processamento de semen equino.',
        'design': (
            '15 garanhoes (5 bons, 5 intermediarios, 5 maus).\n'
            '3 processamentos: Centrifugacao simples (controle) vs SpermFilter vs SLC.\n\n'
            'Congelacao em BotuCrio -> avaliacao pos-descongelacao.\n'
            'Parametros: CASA, membrana, JC-1, morfologia.\n\n'
            'Bonus: IA em subgrupo para taxa de prenhez.\n\n'
            'n = 15 garanhoes x 3 tratamentos = 45 amostras.'
        ),
        'equipment': 'SpermFilter (disponivel), coloide SLC/Androcoll-E (~US$200-400), BotuCrio, CASA, citometria.',
        'cost': 'Moderado (coloide SLC e o item mais caro).',
    },
    {
        'id': 'A7',
        'title': 'Re-suspensao Pos-Refrigeracao Aplicada a Congelacao',
        'literature': [
            ('Segabinazzi et al. (2024)', 'Equine Vet J',
             'Re-suspensao pos-cooling melhora qualidade de bad coolers em semen REFRIGERADO. '
             'Processamento do semen apos o periodo de refrigeracao como alternativa.'),
            ('Papa/Alvarenga — criopreservacao', 'Diversos',
             'Protocolo padrao: diluir -> resfriar -> congelar diretamente. Nenhuma etapa de '
             're-suspensao em diluente fresco antes da congelacao.'),
        ],
        'gap': 'Segabinazzi (2024) mostrou que re-suspensao funciona para REFRIGERACAO. '
               'NINGUEM testou se essa mesma tecnica funciona como etapa ANTES da congelacao — '
               'ou seja, refrigerar -> re-suspender em BotuCrio fresco -> congelar.',
        'differential': 'Aplica uma inovacao recente (2024, do proprio grupo UNESP) em contexto novo. '
                       'Hipotese: re-suspensao remove subprodutos metabolicos e plasma residual '
                       'acumulados, dando um "fresh start" antes do estresse da congelacao.',
        'design': (
            'Bad coolers e intermediarios (n = 12).\n\n'
            '- G1: Protocolo padrao (diluicao -> congelacao direta)\n'
            '- G2: Diluicao -> refrigeracao 2h -> centrifugacao -> re-suspensao em BotuCrio fresco -> congelacao\n'
            '- G3: Diluicao -> refrigeracao 2h -> SpermFilter -> re-suspensao em BotuCrio fresco -> congelacao\n\n'
            'Avaliacao pos-descongelacao: CASA T0/T30, membrana, JC-1.\n\n'
            'n = 12 garanhoes x 3 tratamentos = 36 amostras.'
        ),
        'equipment': 'BotuSemen, BotuCrio, SpermFilter (todos disponIveis), CASA, citometria.',
        'cost': 'Baixo — usa materiais de rotina.',
    },
    {
        'id': 'A8',
        'title': 'Fracionamento do Ejaculado x Qualidade Pos-Congelacao',
        'literature': [
            ('Dell\'Aqua Jr — bad coolers', 'UNESP',
             'Fracionamento do ejaculado listado como ferramenta para bad freezers. '
             'Primeira fracao: maior concentracao, menos plasma seminal.'),
            ('Papa — coleta fracionada', 'UNESP, diversos',
             'Vagina artificial Botucatu permite coleta fracionada. Primeira fracao e rica em '
             'espermatozoides com menor contaminacao de plasma.'),
            ('Fractionated collection for seminal vesiculitis (2020)', 'Theriogenology',
             'Coleta fracionada como ferramenta para garanhoes com vesiculite: primeiras fracoes '
             'tem menor contaminacao de secrecao vesicular.'),
        ],
        'gap': 'NENHUM estudo publicado pela UNESP comparou sistematicamente 1a fracao '
               '(primeiros 2-3 jatos, rica em espermatozoides) vs ejaculado total filtrado na '
               'CONGELABILIDADE com BotuCrio.',
        'differential': 'Teste simples — muda so o manejo de coleta. Hipotese: 1a fracao tem menor '
                       'volume de plasma deletero, melhor para bad freezers cujo plasma e prejudicial.',
        'design': (
            '15 garanhoes.\n'
            'Split: 1a fracao (primeiros 2-3 jatos) vs ejaculado total filtrado.\n\n'
            'Ambos: centrifugacao -> BotuCrio -> congelacao padrao.\n'
            'Avaliacao: CASA, membrana, JC-1, ROS, morfologia.\n\n'
            'n = 15 garanhoes x 3 ejaculados/garanhao = 45 pares.'
        ),
        'equipment': 'VA Botucatu com copo fracionador (disponivel), BotuCrio, CASA, citometria.',
        'cost': 'Baixo — muda so o manejo de coleta.',
    },
    {
        'id': 'A9',
        'title': 'BotuCrio Liofilizado — Prova de Conceito',
        'literature': [
            ('Novel lyophilized extenders (2024)', 'Theriogenology',
             'Diluentes liofilizados mantiveram qualidade equivalente a convencionais. '
             'Eliminam necessidade de cadeia de frio para o diluente. Reconstituicao no momento do uso.'),
            ('Freeze-drying review — future of stallion semen (2024)', 'Vet Sci',
             'Liofilizacao como fronteira. Diluentes liofilizados: boa opcao para logistica. '
             'Semen liofilizado: ainda experimental (motilidade perdida, so ICSI).'),
        ],
        'gap': 'NINGUEM liofilizou o BotuCrio especificamente. Seria inovacao com potencial '
               'comercial enorme — elimina cadeia de frio para transporte do diluente.',
        'differential': 'Prova de conceito que, se funcionar, gera patente e produto comercial '
                       'novo para a Botupharma. Requer colaboracao com Instituto de Biociencias '
                       'ou Quimica (liofilizador).',
        'design': (
            'BotuCrio fresco (controle) vs BotuCrio liofilizado e reconstituido.\n\n'
            '10 garanhoes, 3 ejaculados cada.\n'
            'Congelacao padrao -> avaliacao pos-descongelacao.\n'
            'Parametros: CASA T0/T30, membrana, JC-1, morfologia.\n\n'
            'Testes de estabilidade do liofilizado: 1 semana, 1 mes, 3 meses a 25 graus C.'
        ),
        'equipment': 'BotuCrio, liofilizador (via parceria interdepartamental), CASA, citometria.',
        'cost': 'Moderado (acesso ao liofilizador).',
    },
    {
        'id': 'A10',
        'title': 'Combinacao Piruvato + L-Carnitina no BotuSemen para 72h',
        'literature': [
            ('L-carnitina em BotuSemen', 'Clinical Theriogenology',
             'LC e ALC melhoram motilidade e membrana a 5 graus C ate 48h. '
             'L-carnitina facilita beta-oxidacao -> energia via acetil-CoA -> Krebs.'),
            ('Pena — piruvato', 'Biol Reprod 2021, Theriogenology 2023',
             'Piruvato alimenta Krebs diretamente, regenera NAD+ para glicolise. '
             'Melhora eficiencia mitocondrial.'),
            ('L-carnitina em INRA96 crioprotetor (2025)', 'Theriogenology',
             '5 mM de L-carnitina no meio de congelacao combate crioinjurias em garanhoes Marwari.'),
        ],
        'gap': 'L-carnitina ja foi testada no BotuSemen (publicado). Piruvato nunca. '
               'A COMBINACAO piruvato + L-carnitina NUNCA foi testada em nenhum diluente equino. '
               'Base teorica forte: L-carnitina -> beta-oxidacao -> acetil-CoA; piruvato -> Krebs. '
               'Sinergismo metabolico.',
        'differential': 'Fatorial 2x2 que testa sinergismo metabolico. Se G4 (combinacao) > G2 e G3 '
                       '(individuais), confirma que alimentar duas vias energeticas simultaneamente '
                       'e superior. Elegante e barato.',
        'design': (
            'Fatorial 2x2:\n'
            '- G1: BotuSemen (controle)\n'
            '- G2: BotuSemen + piruvato 5 mM\n'
            '- G3: BotuSemen + L-carnitina 5 mM\n'
            '- G4: BotuSemen + piruvato 5 mM + L-carnitina 5 mM\n\n'
            'Avaliacao: 0h, 24h, 48h, 72h a 5 graus C.\n'
            'Parametros: CASA, JC-1 (mitocondria), membrana, ROS.\n\n'
            'n = 12 garanhoes.'
        ),
        'equipment': 'BotuSemen, piruvato de sodio, L-carnitina (todos baratos), CASA, citometria.',
        'cost': 'Muito baixo.',
    },
]

ch = 4
for proj in diluent_projects:
    write_project(doc, proj, ch)
    ch += 1

# ============================================================
# SYNTHESIS TABLE A
# ============================================================
doc.add_page_break()
doc.add_heading('14. Sintese Parte A — Priorizacao dos Projetos de Diluentes', level=1)

table = doc.add_table(rows=1, cols=6, style='Table Grid')
for i, h in enumerate(['#', 'Projeto', 'Custo', 'Equip. novo?', 'Original.', 'Impacto']):
    table.rows[0].cells[i].text = h
make_header_row(table)

synth_a = [
    ('A10', 'Piruvato + L-carnitina', 'Muito baixo', 'Nao', '★★★★★', 'Alto'),
    ('A1', 'Piruvato no BotuSemen', 'Muito baixo', 'Nao', '★★★★★', 'Alto'),
    ('A8', 'Fracionamento ejaculado', 'Baixo', 'Nao', '★★★★', 'Medio-alto'),
    ('A7', 'Re-suspensao + congelacao', 'Baixo', 'Nao', '★★★★★', 'Alto'),
    ('A4', 'Plasma dose-resposta', 'Baixo', 'Nao', '★★★★★', 'Alto'),
    ('A2', 'Trehalose + prenhez', 'Moderado', 'Nao', '★★★★', 'Muito alto'),
    ('A3', 'DMF bad freezers + IA', 'Moderado', 'Nao', '★★★★', 'Alto'),
    ('A6', 'SpermFilter vs SLC', 'Moderado', 'Coloide', '★★���★', 'Estrategico'),
    ('A5', 'BHT vs BHT-nano', 'Moderado', 'NPs', '★★★★★', 'Alto'),
    ('A9', 'BotuCrio liofilizado', 'Moderado', 'Liofil.', '★★★★', 'Muito alto'),
]
for row_data in synth_a:
    row = table.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        for p in row.cells[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
set_table_style(table)

doc.add_paragraph()
doc.add_paragraph(
    'Projetos A10, A1, A8 e A7 podem ser iniciados imediatamente com materiais de rotina.'
)

doc.add_page_break()

# ============================================================
# PARTE B — CAPACITACAO/FIV
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PARTE B')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Capacitacao Espermatica para FIV Equina')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2E, 0x5E, 0x8E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('10 Projetos')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ============================================================
# 15. ESTADO DA ARTE — FIV
# ============================================================
doc.add_heading('15. Estado da Arte: FIV Equina e Capacitacao (2022-2026)', level=1)

doc.add_heading('15.1 O Problema Historico', level=2)
doc.add_paragraph(
    'A FIV convencional equina fracassou por mais de 30 anos. O espermatozoide equino NAO '
    'realiza efluxo de colesterol in vitro — etapa obrigatoria em bovinos e humanos. Sem efluxo, '
    'a cascata de capacitacao fica incompleta: hiperativacao reversivel, sem reacao acrossomica '
    'funcional, sem penetracao da zona pelucida (Leemans et al., 2019).'
)

doc.add_heading('15.2 Marcos Recentes', level=2)
table = doc.add_table(rows=1, cols=4, style='Table Grid')
for i, h in enumerate(['Ano', 'Grupo', 'Descoberta', 'Resultado']):
    table.rows[0].cells[i].text = h
make_header_row(table)

marcos = [
    ('2022', 'Hinrichs/Felix (Penn Vet)', 'FIV com semen fresco, PHE 22h', '90% fert., 74% blast., 3 potros'),
    ('2025', 'Felix/Hinrichs', 'FIV com semen congelado, PHE 9h', '73% fert., SSD 88,5%'),
    ('2025', 'Martin-Pelaez (UC Davis)', 'FIV congelado = ICSI', '41,1% vs 41,8% blastocistos'),
    ('2026', 'Correa (Univ. Florida)', '1a FIV congelado na Florida', 'Reproducao confirmada'),
    ('Ongoing', 'Leemans (Utrecht)', 'Oviduct-on-chip', 'Fatores oviductais desconhecidos'),
]
for m in marcos:
    add_table_row(table, m)
set_table_style(table)

doc.add_paragraph()

doc.add_heading('15.3 Cascata de Capacitacao Equina — O Bloqueio', level=2)
doc.add_paragraph(
    'Leemans demonstrou que a cascata equina trava no efluxo de colesterol:\n\n'
    'Bicarbonato -> sAC -> cAMP -> PKA -> reorganizacao de membrana -> [BLOQUEIO: sem efluxo '
    'de colesterol] -> fluidez insuficiente -> redistribuicao incompleta de canais ionicos -> '
    'hiperativacao parcial -> fosforilacao parcial de tirosina -> sem reacao acrossomica '
    'funcional -> penetracao da ZP FALHA.\n\n'
    'O protocolo PHE de 22h contorna esse bloqueio por mecanismo ainda nao completamente '
    'elucidado — possivelmente uma via alternativa ativada pelo tempo prolongado.'
)

doc.add_heading('15.4 Dados Criticos — Tempo e Metodo de Separacao', level=2)
doc.add_paragraph(
    'Felix (2025) — Fertilizacao por tempo de pre-incubacao (semen congelado):\n'
    '15 min: 7,1% | 3h: 22,2% | 6h: 38,5% | 9h: 73,3%\n\n'
    'Felix (2025) — Fertilizacao por metodo de separacao:\n'
    'SSD: 88,5% | Swim-up: 38,1% | Centrifugacao coloide: 20,0%'
)

doc.add_page_break()

# ============================================================
# PROJETOS B1-B10 (CAPACITACAO/FIV)
# ============================================================
cap_projects = [
    {
        'id': 'B1',
        'title': 'Reproducao do Protocolo PHE com Semen Congelado em BotuCrio',
        'literature': [
            ('Hinrichs et al. (2022)', 'Biol Reprod, PMID 36084034',
             'Primeiro protocolo FIV equina. Meio TALP + PHE (penicilamina 0,46 mM + hipotaurina '
             '0,23 mM + epinefrina 0,046 mM). Pre-incubacao 20-22h. Fertilizacao ~90%, blastocistos '
             '~74%, 3 potros.'),
            ('Felix et al. (2025)', 'Biol Reprod, PMID 40057974',
             'FIV com semen congelado: 9h pre-incubacao, 73% fertilizacao. SSD = 88,5%. '
             'Fosforilacao de tirosina difere entre fresco e congelado.'),
            ('Martin-Pelaez et al. (2024)', 'Theriogenology',
             'FIV congelado = ICSI em blastocistos (41,1% vs 41,8%).'),
            ('Correa et al. (2026)', 'University of Florida',
             'Primeira FIV com semen congelado na Florida.'),
        ],
        'gap': 'Nenhum laboratorio brasileiro reproduziu o protocolo PHE. Nenhum estudo usou '
               'semen criopreservado em BotuCrio (MF + glicerol + gema 20%).',
        'differential': 'Primeira validacao externa no Brasil. Primeiro teste com BotuCrio — '
                       'posiciona o principal crioprotetor brasileiro como compativel com FIV.',
        'design': (
            'Oocitos: aspiracao de ovarios de abatedouro -> MIV 24-30h -> selecao MII.\n'
            'Semen congelado em BotuCrio -> descongelacao 46 graus C/20s -> lavagem.\n\n'
            'Pre-incubacao em PHE: 4 tempos — 6h, 9h, 12h, 15h.\n'
            'Controle negativo: T0 (sem pre-incubacao).\n'
            'Co-incubacao com oocitos MII por 3h.\n\n'
            'Avaliacao: pronucleos (18-24h), clivagem (48-72h), blastocistos (dia 7-9).\n'
            'Monitoramento: CTC, CASA (hiperativacao), viabilidade em cada tempo.\n\n'
            'n = 30-40 oocitos/grupo, 4-5 repeticoes.'
        ),
        'equipment': 'CASA, citometria, BotuCrio, incubadora CO2, estereomicroscopio. Reagentes PHE: ~R$200-400.',
        'cost': 'Moderado — infraestrutura de cultura embrionaria.',
    },
    {
        'id': 'B2',
        'title': 'Cinetica de Capacitacao x Tipo de Crioprotetor',
        'literature': [
            ('Felix et al. (2025)', 'Biol Reprod, PMID 40057974',
             'Semen congelado capacita em 9h vs 22h (fresco). Padroes de fosforilacao diferem. '
             'Nenhuma subpopulacao hiperpolarizada no congelado.'),
            ('Leemans et al. (2019)', 'Reproduction, PMID 30721132',
             'cAMP basal equino = 10x menor que suino. Ausencia de efluxo de colesterol.'),
            ('Thomas et al. (2006)', 'Reprod Fertil Dev',
             'Criocapacitacao: CTC padrao B sobe de 5,4% para 64,8% apos congelamento.'),
            ('Unraveling glycerol and amides (2024)', 'Theriogenology',
             'Interacao crioprotetor x curva e significativa. MF melhor funcao mitocondrial.'),
        ],
        'gap': 'NENHUM estudo investigou se o tipo de diluente de congelacao afeta a cinetica de '
               'capacitacao para FIV. Se BotuCrio (amida) induz criocapacitacao diferente de INRA96+glicerol, '
               'o tempo otimo de pre-incubacao seria diferente para cada diluente.',
        'differential': 'Primeiro estudo demonstrando que o protocolo de FIV deve ser ajustado conforme '
                       'o crioprotetor. A UNESP tem acesso exclusivo ao BotuCrio para pesquisa.',
        'design': (
            'Split-sample de 10 garanhoes.\n'
            'Congelacao em 3 diluentes: BotuCrio vs INRA96+glicerol vs gema+glicerol convencional.\n\n'
            'Descongelacao -> lavagem -> PHE.\n'
            'Coletas: 0h, 3h, 6h, 9h, 12h.\n\n'
            'Parametros em cada tempo:\n'
            '- Hiperativacao (CASA: ALH >6 um, VCL >150 um/s, LIN <50%)\n'
            '- CTC (F, B, AR)\n'
            '- Fosforilacao de tirosina (anti-pTyr 4G10)\n'
            '- Viabilidade (CFDA/PI)\n'
            '- JC-1\n\n'
            'NAO precisa de oocitos.\n'
            'n = 10 garanhoes x 3 diluentes x 5 tempos = 150 amostras.'
        ),
        'equipment': 'CASA, citometria, anti-pTyr 4G10 (~US$300-500), CTC, JC-1, CFDA/PI.',
        'cost': 'Baixo-moderado.',
    },
    {
        'id': 'B3',
        'title': 'SpermFilter vs Swim-up vs SSD para Selecao Pre-FIV',
        'literature': [
            ('Felix et al. (2025)', 'Biol Reprod',
             'SSD: 88,5% fertilizacao. Swim-up: 38,1%. Centrifugacao coloide: 20,0%.'),
            ('Papa — SpermFilter', 'Diversos',
             'Recuperacao 89,4%, sem dano centrifugo, reutilizavel. NUNCA testado para FIV.'),
            ('Microfluidic selection (2022)', 'J Equine Vet Sci',
             'Chip microfluIdico: DNA superior ao swim-up. NAO testado para FIV convencional.'),
        ],
        'gap': 'SpermFilter NUNCA foi testado como metodo de selecao pre-FIV equina.',
        'differential': 'Se SpermFilter = SSD, a UNESP posiciona produto proprio para mercado de FIV.',
        'design': (
            '10 garanhoes, semen congelado em BotuCrio -> descongelacao.\n'
            '4 metodos: Centrifugacao simples vs Swim-up vs SpermFilter vs SSD.\n\n'
            'Pos-selecao -> PHE 9h -> avaliacao de capacitacao.\n'
            'Se oocitos: co-incubacao e fertilizacao.\n\n'
            'n = 10 garanhoes, 3 repeticoes.'
        ),
        'equipment': 'SpermFilter (disponivel), SSD (~US$50-100/unidade), CASA, citometria, reagentes PHE.',
        'cost': 'Baixo-moderado.',
    },
    {
        'id': 'B4',
        'title': 'Efeito do Plasma Seminal Residual na Capacitacao In Vitro',
        'literature': [
            ('Sieme et al.', 'TiHo Hannover',
             '5-10% plasma: otimo para CRIOPRESERVACAO. Fatores decapacitantes previnem capacitacao prematura.'),
            ('Talluri et al. (2024)', 'Biopreserv Biobanking',
             'Plasma heterOlogo reduz calcio intracelular em espermatozoides equinos criopreservados.'),
            ('Leemans et al. (2016)', 'Reproduction',
             'Albumina + bicarbonato induz aglutinacao e reduz ligacao oviductal >10x.'),
        ],
        'gap': 'NENHUM estudo testou plasma seminal residual na capacitacao para FIV. '
               'Dose otima para criopreservacao (5-10%) pode ser DIFERENTE da dose para capacitacao (possivelmente 0%).',
        'differential': 'Separa dois objetivos historicamente confundidos: plasma para criopreservacao vs capacitacao.',
        'design': (
            '12 garanhoes.\n'
            'Congelacao em BotuCrio com 4 niveis de plasma: 0%, 5%, 10%, 20%.\n\n'
            'Descongelacao -> PHE -> monitoramento: 0h, 3h, 6h, 9h.\n'
            'Parametros: CTC, fosfo-tirosina, hiperativacao (CASA), viabilidade, JC-1.\n\n'
            'n = 12 garanhoes x 4 doses x 4 tempos = 192 amostras.'
        ),
        'equipment': 'SpermFilter ou centrifugacao (disponivel), CASA, citometria, reagentes PHE/CTC.',
        'cost': 'Baixo.',
    },
    {
        'id': 'B5',
        'title': 'Piruvato no Meio de Capacitacao PHE',
        'literature': [
            ('Pena et al. (2021)', 'Biol Reprod, PMID 33864078',
             'Low glucose + high pyruvate: 76,2% vs 61,7% motilidade. Mitocondrias ativas: 51,1% vs 24,1%.'),
            ('Pena et al. (2023)', 'Theriogenology, PMID 38029686',
             'Piruvato melhora funcao em meios com alta glicose. Regenera NAD+ via conversao a lactato.'),
            ('Protocolo PHE (Hinrichs 2022)', 'Biol Reprod',
             'TALP padrao — nenhuma otimizacao de glicose/piruvato para FIV.'),
        ],
        'gap': 'NINGUEM otimizou razao glicose/piruvato no meio de CAPACITACAO para FIV equina.',
        'differential': 'Cruza linha metabolica (Pena) com FIV (Hinrichs). Duas linhas consolidadas NUNCA integradas.',
        'design': (
            '4 meios de capacitacao:\n'
            '- G1: PHE padrao (TALP — controle)\n'
            '- G2: PHE com piruvato 2x\n'
            '- G3: PHE com glicose 50% + piruvato 2x\n'
            '- G4: PHE sem glicose + piruvato 3x\n\n'
            'Semen congelado BotuCrio -> selecao -> 9h.\n'
            'Parametros: JC-1, CASA (hiperativacao), CTC, viabilidade, ROS (DHE).\n\n'
            'n = 10 garanhoes.'
        ),
        'equipment': 'Piruvato de sodio, glicose (<R$100). CASA, citometria.',
        'cost': 'Muito baixo.',
    },
    {
        'id': 'B6',
        'title': 'Termotaxia como Selecao de Espermatozoides Capacitados Pre-FIV',
        'literature': [
            ('Ruiz-Diaz et al. (2020)', 'Animals, PMID 32825582',
             'D-penicilamina + HTF: capacitacao equina com selecao por termotaxia. '
             'Fracao selecionada: mais fosfo-tirosina, menos fragmentacao de DNA.'),
            ('Perez-Cerezales (2022)', 'J Anim Sci Biotechnol',
             'Termotaxia bovina: melhor DNA, melhor ICSI.'),
        ],
        'gap': 'NINGUEM aplicou termotaxia como selecao pre-FIV convencional equina.',
        'differential': 'Selecao FUNCIONAL (so capacitados migram). Elimina problema de heterogeneidade temporal.',
        'design': (
            'Semen BotuCrio -> SpermFilter -> PHE 9h.\n\n'
            '- G1: Co-incubacao direta (controle Hinrichs)\n'
            '- G2: Termotaxia (37->39 graus C, 30 min) -> fracao migrada -> co-incubacao\n\n'
            'Avaliar fracao: fosfo-tirosina, CTC, DNA (comet).\n'
            'Endpoints: fertilizacao, clivagem, blastocistos.\n\n'
            'n = 10 garanhoes.'
        ),
        'equipment': 'Camara de termotaxia (<R$500 — construcao simples), CASA, citometria, cultura embrionaria.',
        'cost': 'Baixo.',
    },
    {
        'id': 'B7',
        'title': 'pH Alcalino (7,9) no Meio de Capacitacao PHE',
        'literature': [
            ('Leemans et al. (2015)', 'Reproduction, PMID 26078083',
             'Fluido folicular pH 7,9 induz hipermotilidade Ca2+-dependente. So em espermatozoides '
             'ligados a explantes oviductais.'),
            ('Gonzalez-Fernandez (2013)', 'Biol Reprod',
             'Ligacao ao oviduto + pH elevado induz fosforilacao de tirosina. pH intracelular sobe gradualmente.'),
            ('Leemans (2019)', 'Biol Reprod',
             'Procaina (base fraca): hiperativacao robusta mas SEM fertilizacao. pH-dependente, parcialmente independente de cAMP.'),
        ],
        'gap': 'PHE de Hinrichs usa pH ~7,4. Leemans demonstrou pH 7,9 como gatilho critico. '
               'NINGUEM testou pH elevado no meio PHE de FIV.',
        'differential': 'Integra PHE (empirico) com pH alcalino (mecanistico). Teste simples, alto potencial.',
        'design': (
            '4 pHs no PHE: 7,2 / 7,4 (controle) / 7,6 / 7,9.\n'
            'Tamponamento: HEPES 25 mM + bicarbonato 15 mM.\n\n'
            'Semen BotuCrio -> selecao -> PHE em cada pH -> 9h.\n'
            'Parametros: hiperativacao (CASA), CTC, fosfo-tirosina, viabilidade, JC-1.\n'
            'Se oocitos: co-incubacao e fertilizacao.\n\n'
            'n = 10 garanhoes.'
        ),
        'equipment': 'NaOH, HCl, HEPES, pHmetro (tudo disponivel). CASA, citometria.',
        'cost': 'Muito baixo.',
    },
    {
        'id': 'B8',
        'title': 'Bad Freezers na FIV: Criocapacitacao como Vantagem Paradoxal?',
        'literature': [
            ('Thomas et al. (2006)', 'Reprod Fertil Dev',
             'Criocapacitacao: CTC B sobe de 5,4% para 64,8%. AR de 1% para 32,8%.'),
            ('Felix et al. (2025)', 'Biol Reprod',
             'Semen congelado: 9h vs 22h. Criocapacitacao avanca ponto de partida.'),
            ('Sieme — 30% bad freezers', 'TiHo Hannover',
             'Bad freezers sofrem MAIS criocapacitacao. Associada a subfertilidade em IA.'),
            ('Cellular consequences (2024)', 'J Equine Vet Sci',
             'Criocapacitacao: negativa para IA. NUNCA testada no contexto de FIV.'),
        ],
        'gap': 'Criocapacitacao sempre vista como NEGATIVA. Para FIV, capacitacao e o OBJETIVO. '
               'Se bad freezers sofrem mais criocapacitacao, poderiam paradoxalmente ter melhor FIV.',
        'differential': 'Hipotese contraintuitiva que muda paradigma. Se confirmada, resgata genetica '
                       'de garanhoes considerados problematicos. UNESP tem garanhoes classificados.',
        'design': (
            '15 garanhoes: 5 bons, 5 intermediarios, 5 maus congeladores.\n\n'
            'BotuCrio -> descongelacao -> PHE -> monitoramento: 0, 3, 6, 9, 12h.\n'
            'Parametros: CTC, fosfo-tirosina, hiperativacao, viabilidade, JC-1.\n\n'
            'Endpoint: tempo otimo de cada CLASSE.\n'
            'Se oocitos: FIV no tempo otimo de cada classe.\n\n'
            'n = 15 garanhoes, 3 ejaculados cada.'
        ),
        'equipment': 'CASA, citometria, BotuCrio, reagentes PHE/CTC. Garanhoes classificados (disponivel).',
        'cost': 'Baixo.',
    },
    {
        'id': 'B9',
        'title': 'Antioxidante no Meio de Capacitacao: Proteger sem Inibir',
        'literature': [
            ('Pena (2019)', 'Antioxidants, PMID 31752408',
             'ROS sao essenciais para capacitacao. Eustress vs distress. "Corda bamba redox".'),
            ('Gibb', 'Diversos',
             'Curva em U invertido. Garanhoes mais ferteis = MAIS ROS (paradoxo).'),
            ('Catalase in stallion sperm (2025)', 'Antioxidants',
             'Equinos: 14x mais resistentes a dano oxidativo que humanos (IC50 391,6 vs 27,3 uM).'),
            ('ROS e quimiotaxia (2020)', 'Reproduction',
             'ROS participam da quimiotaxia espermatica equina via NOX -> cAMP-PKA.'),
        ],
        'gap': 'Antioxidantes extensivamente estudados em DILUENTES. NINGUEM testou no meio de '
               'CAPACITACAO para FIV. O paradoxo: proteger sem inibir capacitacao.',
        'differential': 'Primeira curva dose-resposta de antioxidante no contexto de capacitacao para FIV. '
                       'Resolve: adicionar ou nao antioxidante ao PHE?',
        'design': (
            'PHE + 4 doses de MitoTEMPO: 0 (controle), 1 uM, 10 uM, 100 uM.\n\n'
            'Semen BotuCrio -> selecao -> PHE + MitoTEMPO -> 9h.\n\n'
            'Parametros pro-capacitacao: CTC B, fosfo-tirosina, hiperativacao.\n'
            'Parametros pro-viabilidade: membrana, JC-1, DNA.\n'
            'Parametros mecanisticos: ROS (DHE), 4-HNE.\n\n'
            'n = 10 garanhoes.'
        ),
        'equipment': 'MitoTEMPO (~US$150-200). CASA, citometria.',
        'cost': 'Baixo-moderado.',
    },
    {
        'id': 'B10',
        'title': 'Metabolitos Oviductais no Meio PHE',
        'literature': [
            ('Lamy et al. (2020)', 'Theriogenology, PMID 31835098',
             'Metabolomica oviductal equina: lactato, mioinositol, creatina, alanina, carnitina. '
             'Fumarato e glicina diferem entre pre e pos-ovulatorio.'),
            ('Gonzalez-Fernandez (2022)', 'J Equine Vet Sci, PMID 35077851',
             'Metabolitos oviductais NAO modificaram capacitacao — mas usaram protocolo diferente '
             'do PHE (sem pre-incubacao prolongada, concentracoes diferentes).'),
            ('Leemans (2022)', 'Biol Reprod',
             'EOEC em Transwell: celulas secretam fatores que melhoram capacitacao. Identidade desconhecida.'),
        ],
        'gap': 'Gonzalez-Fernandez (2022) concluiu que metabolitos nao funcionam — mas com '
               'protocolo DIFERENTE do PHE. NINGUEM testou no contexto PHE de 9h.',
        'differential': 'Mimetizar microambiente oviductal SEM co-cultura. Simplificacao que viabiliza '
                       'FIV em labs sem cultura celular. Se funcionar, cria "PHE oviductal".',
        'design': (
            '5 grupos:\n'
            '- G1: PHE padrao (controle)\n'
            '- G2: PHE + lactato 25 mM\n'
            '- G3: PHE + mioinositol 5 mM + creatina 1 mM\n'
            '- G4: PHE + L-carnitina 5 mM\n'
            '- G5: PHE + coquetel completo\n\n'
            'Semen BotuCrio -> selecao -> 9h.\n'
            'Parametros: CTC, fosfo-tirosina, hiperativacao, viabilidade, JC-1.\n'
            'Se oocitos: co-incubacao e fertilizacao.\n\n'
            'n = 10 garanhoes. Todos reagentes baratos (<R$300 total).'
        ),
        'equipment': 'Lactato, mioinositol, creatina, L-carnitina, alanina (rotina). CASA, citometria.',
        'cost': 'Muito baixo.',
    },
]

ch = 16
for proj in cap_projects:
    write_project(doc, proj, ch)
    ch += 1

# ============================================================
# SYNTHESIS TABLE B
# ============================================================
doc.add_page_break()
doc.add_heading('26. Sintese Parte B — Priorizacao dos Projetos de Capacitacao/FIV', level=1)

table = doc.add_table(rows=1, cols=6, style='Table Grid')
for i, h in enumerate(['#', 'Projeto', 'Custo', 'Oocitos?', 'Original.', 'Impacto']):
    table.rows[0].cells[i].text = h
make_header_row(table)

synth_b = [
    ('B5', 'Piruvato no PHE', 'Muito baixo', 'Opcional', '★★★★★', 'Alto'),
    ('B7', 'pH 7,9 no PHE', 'Muito baixo', 'Opcional', '★★★★★', 'Muito alto'),
    ('B10', 'Metabolitos oviductais', 'Muito baixo', 'Opcional', '★★★★', 'Alto'),
    ('B2', 'Cinetica x crioprotetor', 'Baixo-mod', 'Nao', '★★★★★', 'Muito alto'),
    ('B4', 'Plasma x capacitacao', 'Baixo', 'Opcional', '★★★★', 'Alto'),
    ('B8', 'Bad freezers FIV', 'Baixo', 'Opcional', '★★★★★', 'Transformador'),
    ('B9', 'Antioxidante no PHE', 'Baixo-mod', 'Nao', '★★★★★', 'Fundamental'),
    ('B1', 'PHE + BotuCrio', 'Moderado', 'Sim', '★★★', 'Estrategico'),
    ('B3', 'SpermFilter pre-FIV', 'Baixo-mod', 'Opcional', '★★★★', 'Estrategico'),
    ('B6', 'Termotaxia pre-FIV', 'Baixo', 'Sim', '★★★★★', 'Alto'),
]
for row_data in synth_b:
    row = table.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        for p in row.cells[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
set_table_style(table)

doc.add_page_break()

# ============================================================
# 27. PROPOSTA DE INTEGRACAO
# ============================================================
doc.add_heading('27. Proposta de Integracao: Duas Teses de Doutorado', level=1)

doc.add_heading('Tese 1 — Diluentes', level=2)
p = doc.add_paragraph()
run = p.add_run('"Otimizacao de diluentes para criopreservacao de semen equino: '
                'suplementacao metabolica, modulacao de plasma seminal e novas tecnologias"')
run.italic = True

chapters_a = [
    ('Capitulo 1 (A10)', 'Piruvato + L-carnitina: sinergismo metabolico no BotuSemen'),
    ('Capitulo 2 (A4)', 'Plasma seminal dose-resposta x classificacao de garanhoes no BotuCrio'),
    ('Capitulo 3 (A7+A8)', 'Re-suspensao pre-congelacao + fracionamento do ejaculado em bad freezers'),
    ('Capitulo 4 (A2)', 'Trehalose no BotuCrio: validacao in vivo com taxa de prenhez'),
]
for ch_t, ch_d in chapters_a:
    p = doc.add_paragraph()
    run = p.add_run(f'{ch_t}: ')
    run.bold = True
    p.add_run(ch_d)

doc.add_paragraph()

doc.add_heading('Tese 2 — Capacitacao/FIV', level=2)
p = doc.add_paragraph()
run = p.add_run('"Otimizacao da capacitacao in vitro de espermatozoides equinos criopreservados '
                'em BotuCrio para fertilizacao in vitro convencional"')
run.italic = True

chapters_b = [
    ('Capitulo 1 (B2)', 'Cinetica de capacitacao em funcao do tipo de crioprotetor'),
    ('Capitulo 2 (B5+B7)', 'Otimizacao metabolica (piruvato) e de pH (7,9) no meio PHE'),
    ('Capitulo 3 (B8)', 'Bad freezers na FIV: criocapacitacao como vantagem paradoxal'),
    ('Capitulo 4 (B1+B3)', 'Validacao integrada: BotuCrio + SpermFilter + PHE otimizado -> FIV'),
]
for ch_t, ch_d in chapters_b:
    p = doc.add_paragraph()
    run = p.add_run(f'{ch_t}: ')
    run.bold = True
    p.add_run(ch_d)

doc.add_paragraph()
doc.add_paragraph(
    'As duas teses sao complementares: a Tese 1 otimiza a criopreservacao do semen, e a Tese 2 '
    'usa esse semen otimizado para FIV. Juntas, cobrem a cadeia completa desde a coleta ate a '
    'producao de embrioes in vitro. Estimativa total: 8 capitulos, 8-10 publicacoes, impacto '
    'internacional com posicionamento de produtos UNESP/Botupharma.'
)

doc.add_page_break()

# ============================================================
# 28. REFERENCIAS
# ============================================================
doc.add_heading('28. Referencias Bibliograficas', level=1)

references = [
    'Hinrichs K, Choi YH, Love CC, Chung YG, Varner DD. Successful in vitro fertilization in the horse: production of blastocysts and birth of foals after prolonged sperm incubation for capacitation. Biol Reprod. 2022;107(6):1320-1329. PMID: 36084034.',
    'Felix MR, Loux SC, Esteller-Vico A, Hinrichs K. Equine in vitro fertilization with frozen-thawed semen is associated with shortened pre-incubation time and modified capacitation-related changes. Biol Reprod. 2025;112(5):867-878. PMID: 40057974.',
    'Martin-Pelaez P, et al. IVF with frozen-thawed sperm after prolonged capacitation yields comparable results to ICSI in horses: A morphokinetics study. Theriogenology. 2024.',
    'Leemans B, Stout TAE, et al. Update on mammalian sperm capacitation: how much does the horse differ from other species? Reproduction. 2019;157(5):R181-R197. PMID: 30721132.',
    'Leemans B, et al. An alkaline follicular fluid fraction induces capacitation and limited release of oviduct epithelium-bound stallion sperm. Reproduction. 2015;150(3):193-208. PMID: 26078083.',
    'Leemans B, et al. Developing a reproducible protocol for culturing functional confluent monolayers of differentiated equine oviduct epithelial cells. Biol Reprod. 2022;106(4):710-726.',
    'Leemans B, et al. pH-dependent effects of procaine on equine gamete activation. Biol Reprod. 2019;101(5):1056-1074.',
    'Pena FJ, et al. Low glucose and high pyruvate reduce the production of 2-oxoaldehydes, improving mitochondrial efficiency, redox regulation, and stallion sperm function. Biol Reprod. 2021;105(2):519-532. PMID: 33864078.',
    'Pena FJ, et al. Pyruvate enhances stallion sperm function in high glucose media improving overall metabolic efficiency. Theriogenology. 2023. PMID: 38029686.',
    'Pena FJ, et al. Redox Regulation and Oxidative Stress: The Particular Case of the Stallion Spermatozoa. Antioxidants. 2019;8(11):567. PMID: 31752408.',
    'Gibb Z, Aitken RJ. The impact of sperm metabolism during in vitro storage: the stallion as a model. BioMed Res Int. 2016.',
    'Lamy J, et al. Stage-specific metabolomic changes in equine oviductal fluid. Theriogenology. 2020;143:140-151. PMID: 31835098.',
    'Gonzalez-Fernandez L, et al. Selected Metabolites Found in Equine Oviductal Fluid do not Modify the Parameters Associated to Capacitation. J Equine Vet Sci. 2022;111:103875. PMID: 35077851.',
    'Ruiz-Diaz S, et al. D-Penicillamine during In Vitro Capacitation of Stallion Spermatozoa Prolongs Hyperactive-Like Motility and Allows Sperm Selection by Thermotaxis. Animals. 2020;10(9):1467. PMID: 32825582.',
    'Maitan PP, et al. A stallion spermatozoon\'s journey through the mare\'s genital tract: In vivo and in vitro aspects of sperm capacitation. Anim Reprod Sci. 2022;236:106932.',
    'Thomas AD, et al. Capacitation-like changes in equine spermatozoa throughout the cryopreservation process. Reprod Fertil Dev. 2006;14(4):225-233.',
    'Sieme H, et al. Advances in Stallion Semen Cryopreservation. Vet Clin North Am Equine Pract. 2016;32(3):521-530.',
    'Papa FO, Alvarenga MA. Advances in Stallion Semen Cryopreservation. Vet Clin North Am Equine Pract. 2016;32(3):521-530.',
    'Segabinazzi L, et al. Post-cooling semen processing and sperm re-suspension as an alternative method to circumvent poor semen cooling in stallions. Equine Vet J. 2024.',
    'Efficacy of butylated hydroxytoluene nanoparticles in enhancing the quality and preservation of stallion chilled semen. Vet Res Commun. 2025.',
    'Selenium nanoparticles-supplemented INRA96 extender on Turkmen stallion sperm quality. Theriogenology. 2024.',
    'Evaluation of a Chemically Defined, Long-Term Extender (Beyond) for Liquid Storage of Stallion Semen. Reprod Domest Anim. 2025.',
    'Freezability and Fertility Rates of Stallion Semen Supplemented With Trehalose in Lactose Extender. J Equine Vet Sci. 2023.',
    'Innovative Approaches to Avoid Antibiotic Use in Equine Semen Cryopreservation. Animals. 2025;15(10):1368.',
    'Novel lyophilized extenders maintain stallion semen characteristics during freezing. Theriogenology. 2024.',
    'Enhancing Stallion Semen Cryopreservation: Selected Antioxidant Extracts and Sperm Freezability. Antioxidants. 2025;14(11):1363.',
    'Unraveling the effect of glycerol and amides and different cooling rates on post-thawed sperm parameters. Theriogenology. 2024.',
    'Talluri TR, et al. Heterologous Seminal Plasma Reduces Intracellular Calcium and Sperm Viability of Cryopreserved Stallion Spermatozoa. Biopreserv Biobanking. 2024.',
    'Catalase in Unexpected Places: Revisiting H2O2 Detoxification Pathways in Stallion Spermatozoa. Antioxidants. 2025;14(6):718.',
    'ROS are involved in the signaling of equine sperm chemotaxis. Reproduction. 2020;159(4).',
    'L-carnitine and acetyl-L-carnitine enhance stallion sperm quality during semen storage at 5 degrees C. Clin Theriogenology. 2024.',
    'Challenges and Enhancing Strategies of Equine Semen Preservation: Nutritional and Genetic Perspectives. Vet Sci. 2025;12(9):807.',
    'Correa L, et al. Equine IVF milestone. University of Florida. 2026.',
    'Conserved Mechanism of Bicarbonate-Induced Sensitization of CatSper Channels. Front Cell Dev Biol. 2021.',
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'[{i}] ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(ref)
    run.font.size = Pt(9)

# ============================================================
# SAVE
# ============================================================
output_path = '/Users/gabrielpetelinkar/AIOSORIZA/outputs/20_Projetos_Pesquisa_Reproducao_Equina_UNESP.docx'
doc.save(output_path)
print(f'Documento salvo em: {output_path}')
print(f'Tamanho estimado: ~80+ paginas')
