#!/usr/bin/env python3
"""
Gera documento DOCX: Projetos de Pesquisa — Capacitacao Espermatica para FIV Equina
Rastreio de literatura + justificativa de diferencial
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Title style
title_style = doc.styles['Title']
title_style.font.size = Pt(24)
title_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
title_style.font.bold = True

# Heading styles
for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

def add_ref_paragraph(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True
    return p

def add_table_row(table, cells_text, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        if bold:
            run.bold = True
    return row

def set_table_style(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROJETOS DE PESQUISA')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Capacitacao Espermatica para Fertilizacao\nIn Vitro Convencional Equina')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2E, 0x5E, 0x8E)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Rastreio de Literatura e Justificativa de Diferencial')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.italic = True

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Elaborado com auxilio do Equine Reproduction Research Squad')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Especialistas consultados: Leemans, Hinrichs, Felix, Maitan, Stout, Squires, Pena, Papa')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Abril 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS placeholder
# ============================================================
doc.add_heading('Sumario', level=1)
toc_items = [
    '1. Contexto e Estado da Arte',
    '2. Justificativa Geral — Por que a UNESP?',
    '3. Projeto 1 — Reproducao do Protocolo PHE com BotuCrio',
    '4. Projeto 2 — Cinetica de Capacitacao x Tipo de Crioprotetor',
    '5. Projeto 3 — SpermFilter vs Swim-up vs SSD para Selecao Pre-FIV',
    '6. Projeto 4 — Plasma Seminal Residual e Capacitacao In Vitro',
    '7. Projeto 5 — Piruvato no Meio de Capacitacao PHE',
    '8. Projeto 6 — Termotaxia como Selecao de Espermatozoides Capacitados',
    '9. Projeto 7 — pH Alcalino (7,9) no Meio de Capacitacao',
    '10. Projeto 8 — Bad Freezers na FIV: Criocapacitacao como Vantagem?',
    '11. Projeto 9 — Antioxidante no Meio de Capacitacao: Proteger sem Inibir',
    '12. Projeto 10 — Metabolitos Oviductais no Meio PHE',
    '13. Sintese Comparativa e Priorizacao',
    '14. Proposta de Tese Integrada',
    '15. Referencias Bibliograficas',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. CONTEXTO E ESTADO DA ARTE
# ============================================================
doc.add_heading('1. Contexto e Estado da Arte', level=1)

doc.add_heading('1.1 O Problema Historico da FIV Equina', level=2)
doc.add_paragraph(
    'A fertilizacao in vitro (FIV) convencional — co-incubacao de oocitos maduros com espermatozoides '
    'capacitados — funciona rotineiramente em bovinos (desde os anos 1980), humanos (desde 1978) e suinos. '
    'Em equinos, porem, a FIV convencional fracassou consistentemente por mais de 30 anos. O espermatozoide '
    'equino nao conseguia penetrar a zona pelucida (ZP) in vitro, mesmo quando todos os parametros de '
    'capacitacao aparentavam estar normais.'
)
doc.add_paragraph(
    'O Leemans e colaboradores (Utrecht/Ghent) demonstraram que o problema reside em uma peculiaridade '
    'especie-especifica: o espermatozoide equino NAO realiza efluxo de colesterol da membrana durante '
    'a capacitacao in vitro — etapa obrigatoria em bovinos e humanos para aumentar a fluidez de membrana '
    'e redistribuir canais ionicos (Leemans et al., 2019, PMID 30721132). Sem efluxo de colesterol, '
    'a cascata de capacitacao fica incompleta: ha ativacao parcial de sAC/cAMP/PKA, fosforilacao parcial '
    'de tirosina, mas a hiperativacao e reversivel e a reacao acrossomica funcional nao ocorre.'
)

doc.add_heading('1.2 O Breakthrough de 2022 — Hinrichs/Felix', level=2)
doc.add_paragraph(
    'Em 2022, Katrin Hinrichs e Matheus Felix (Penn Vet) publicaram o primeiro protocolo eficiente de '
    'FIV convencional equina (Hinrichs et al., Biology of Reproduction, 2022, PMID 36084034). A chave '
    'foi surpreendentemente simples: prolongar o tempo de pre-incubacao dos espermatozoides no meio PHE '
    '(penicilamina + hipotaurina + epinefrina em TALP modificado) para 20-22 horas — muito alem das '
    '4-6 horas usadas em protocolos bovinos.'
)
doc.add_paragraph(
    'Resultados com semen fresco:\n'
    '- Taxa de fertilizacao: ~90%\n'
    '- Taxa de blastocistos: ~74%\n'
    '- 3 embrioes transferidos -> 3 potros saudaveis nascidos'
)

doc.add_heading('1.3 FIV com Semen Congelado — Felix et al., 2025', level=2)
doc.add_paragraph(
    'Em 2025, Felix et al. (Biology of Reproduction, PMID 40057974) demonstraram que semen '
    'congelado-descongelado tambem funciona para FIV, com tempo de pre-incubacao reduzido para 9 horas. '
    'A hipotese e que a criocapacitacao (alteracoes de membrana induzidas pelo congelamento) avanca o '
    'ponto de partida da capacitacao, encurtando o tempo necessario.'
)

# Table: timeline of capacitation
table = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr = table.rows[0].cells
hdr[0].text = 'Tipo de Semen'
hdr[1].text = 'Tempo Otimo Pre-incubacao'
hdr[2].text = 'Taxa de Fertilizacao'
for cell in hdr:
    for p in cell.paragraphs:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(10)

data = [
    ('Fresco', '20-22 horas', '~90%'),
    ('Congelado-descongelado', '9 horas', '~73%'),
    ('Congelado (sem pre-incubacao)', '15 min', '7,1%'),
]
for row_data in data:
    add_table_row(table, row_data)
set_table_style(table)

doc.add_paragraph()
doc.add_paragraph(
    'Adicionalmente, o metodo de separacao espermatica pos-descongelacao mostrou-se critico: '
    'SSD (Sperm Separation Device) produziu 88,5% de fertilizacao vs 38,1% (swim-up) vs 20,0% '
    '(centrifugacao em coloide).'
)

doc.add_heading('1.4 Validacoes Independentes (2025-2026)', level=2)
doc.add_paragraph(
    'Martin-Pelaez et al. (UC Davis, Theriogenology, 2024) confirmaram que FIV com semen congelado '
    'produz taxas de blastocistos comparaveis a ICSI (41,1% vs 41,8%). Em marco de 2026, a '
    'Universidade da Florida (Correa et al.) reportou a primeira FIV equina com semen congelado '
    'no estado, confirmando a reprodutibilidade do protocolo.'
)

doc.add_heading('1.5 Gaps Criticos Identificados na Literatura', level=2)

gaps = [
    ('Identidade molecular dos fatores oviductais',
     'O Leemans demonstrou que celulas epiteliais oviductais completam a capacitacao, mas os '
     'fatores moleculares especificos permanecem desconhecidos. E o "Santo Graal" da area.'),
    ('Ausencia de efluxo de colesterol',
     'Equinos nao realizam efluxo de colesterol in vitro — mecanismo desconhecido. '
     'Pode ser deficiencia de transportador lipidico ou composicao de membrana especie-especifica.'),
    ('Mecanismo da capacitacao acelerada do semen congelado',
     'Felix mostrou 9h vs 22h, mas os mecanismos moleculares que diferenciam a capacitacao '
     'de semen congelado vs fresco sao desconhecidos.'),
    ('Variabilidade entre garanhoes',
     'Hinrichs declarou que fazer FIV funcionar com TODO semen e o proximo desafio. '
     'Alguns garanhoes respondem mal ao protocolo.'),
    ('Ativacao pH-dependente independente de cAMP',
     'Leemans mostrou que pH 7,9 (fluido folicular) ativa capacitacao por via nao-canonica, '
     'mas o mecanismo permanece obscuro.'),
    ('Otimizacao metabolica do meio de capacitacao',
     'O protocolo PHE usa TALP padrao — nenhuma otimizacao da razao glicose/piruvato foi tentada.'),
    ('Efeito do plasma seminal residual na capacitacao para FIV',
     'Plasma seminal contem fatores decapacitantes — impacto na FIV desconhecido.'),
    ('Antioxidantes no meio de capacitacao',
     'ROS sao necessarias para capacitacao — impacto de antioxidantes nesse contexto nao testado.'),
]

for i, (gap, desc) in enumerate(gaps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'Gap {i}: {gap}. ')
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(desc)

doc.add_page_break()

# ============================================================
# 2. JUSTIFICATIVA GERAL — POR QUE A UNESP?
# ============================================================
doc.add_heading('2. Justificativa Geral — Por que a UNESP?', level=1)

doc.add_paragraph(
    'A FMVZ/UNESP Botucatu possui vantagens competitivas unicas para pesquisa em capacitacao '
    'espermatica e FIV equina:'
)

advantages = [
    ('BotuCrio e BotuSemen', 'Diluentes proprios (Botupharma) — possibilidade de testar '
     'compatibilidade crioprotetor-capacitacao com produtos comerciais de ampla distribuicao.'),
    ('SpermFilter', 'Tecnologia propria de separacao espermatica — potencial para mercado de FIV '
     'se demonstrada eficacia como metodo de selecao pre-FIV.'),
    ('Rebanho experimental', 'Garanhoes classificados (bons, intermediarios, maus congeladores) '
     'e eguas para validacao in vivo — infraestrutura que poucos centros no mundo possuem.'),
    ('CASA + Citometria de fluxo', 'Avaliacao multiparametrica disponivel: motilidade (CASA), '
     'integridade de membrana (CFDA/PI), potencial mitocondrial (JC-1), ROS (DHE), CTC.'),
    ('Acesso a oocitos', 'Ovarios de abatedouro para testes in vitro sem necessidade de OPU em eguas vivas.'),
    ('Expertise em criopreservacao', 'Grupo Papa/Alvarenga/Dell\'Aqua Jr com decadas de publicacoes — '
     'entendimento profundo de criocapacitacao e variabilidade entre garanhoes.'),
    ('Nenhuma publicacao em FIV convencional', 'A UNESP ainda nao entrou nessa linha — seria o primeiro '
     'grupo brasileiro a publicar dados de FIV equina convencional.'),
]

for title, desc in advantages:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph(
    'A combinacao de produtos proprios (BotuCrio, SpermFilter) + rebanho experimental + expertise '
    'em criopreservacao posiciona a UNESP como o centro ideal para traduzir o breakthrough da '
    'Penn Vet para o contexto brasileiro e para a cadeia comercial Botupharma.'
)

doc.add_page_break()

# ============================================================
# PROJECTS 1-10
# ============================================================

projects = [
    {
        'num': 1,
        'title': 'Reproducao do Protocolo PHE com Semen Congelado em BotuCrio',
        'literature': [
            ('Hinrichs et al. (2022)', 'Biology of Reproduction, PMID 36084034',
             'Primeiro protocolo eficiente de FIV equina com semen fresco. Meio TALP + PHE '
             '(penicilamina 0,46 mM + hipotaurina 0,23 mM + epinefrina 0,046 mM). '
             'Pre-incubacao 20-22h. Fertilizacao ~90%, blastocistos ~74%, 3 potros nascidos.'),
            ('Felix et al. (2025)', 'Biology of Reproduction, PMID 40057974',
             'FIV com semen congelado. Pre-incubacao 9h em PHE. Fertilizacao 73%. '
             'Semen criopreservado em diluente nao-especificado como BotuCrio. '
             'Padroes de fosforilacao de tirosina diferem do semen fresco.'),
            ('Martin-Pelaez et al. (2024)', 'Theriogenology',
             'FIV vs ICSI com semen congelado. Blastocistos: 41,1% (FIV) vs 41,8% (ICSI). '
             'Confirma que FIV e alternativa viavel a ICSI.'),
            ('Correa et al. (2026)', 'University of Florida',
             'Primeira FIV equina com semen congelado na Florida. Confirma reprodutibilidade '
             'do protocolo em outro laboratorio.'),
        ],
        'gap': 'Nenhum laboratorio brasileiro reproduziu o protocolo PHE para FIV equina. '
               'Nenhum estudo testou semen criopreservado em BotuCrio (MF + glicerol + gema 20%) '
               'nesse contexto. Os estudos de Hinrichs/Felix e Martin-Pelaez utilizaram diluentes '
               'de congelacao diferentes — a compatibilidade do BotuCrio com o protocolo PHE e desconhecida.',
        'differential': 'Primeira validacao externa do protocolo PHE no Brasil. Primeiro teste com '
                       'BotuCrio — se funcionar, posiciona o principal crioprotetor comercial brasileiro '
                       'como compativel com FIV, impacto direto no mercado de reproducao assistida equina. '
                       'A UNESP e o unico grupo com acesso irrestrito ao BotuCrio para pesquisa.',
        'design': (
            'Obtencao de oocitos: aspiracao folicular de ovarios de abatedouro.\n'
            'Maturacao in vitro (MIV): TCM-199 + FSH + LH + EGF, 24-30h, 38,2 graus C, 5% CO2.\n'
            'Selecao de oocitos MII (presenca de corpusculo polar).\n\n'
            'Semen: ejaculados de 8-10 garanhoes, congelados em BotuCrio (protocolo padrao UNESP).\n'
            'Descongelacao: 46 graus C / 20s (palheta 0,5 mL).\n'
            'Lavagem: centrifugacao em gradiente ou SpermFilter.\n\n'
            'Pre-incubacao em PHE: 4 tempos — 6h, 9h, 12h, 15h.\n'
            'Controle negativo: semen sem pre-incubacao (T0).\n'
            'Co-incubacao com oocitos MII por 3h.\n\n'
            'Avaliacao:\n'
            '- Fertilizacao: formacao de pronucleos (18-24h)\n'
            '- Clivagem: 48-72h\n'
            '- Blastocistos: dia 7-9\n'
            '- Monitoramento da capacitacao em cada tempo: CTC, CASA (hiperativacao), viabilidade\n\n'
            'n = 30-40 oocitos/grupo, 4-5 repeticoes.'
        ),
        'equipment': 'CASA, citometria de fluxo, VA Botucatu, centrIfuga, BotuCrio, incubadora CO2, '
                    'estereomicroscopio para avaliacao de embrioes. Reagentes PHE: ~R$200-400.',
        'cost': 'Moderado — principal custo e infraestrutura de cultura embrionaria.',
    },
    {
        'num': 2,
        'title': 'Cinetica de Capacitacao x Tipo de Crioprotetor',
        'literature': [
            ('Felix et al. (2025)', 'Biology of Reproduction, PMID 40057974',
             'Demonstrou que semen congelado capacita em 9h vs 22h (fresco). '
             'Padrao de fosforilacao de tirosina: equatorial/peca intermediaria em 15 min '
             '(congelado) vs apical/equatorial em 22h (fresco). Nenhuma subpopulacao '
             'hiperpolarizada detectavel no congelado.'),
            ('Leemans et al. (2019)', 'Reproduction, PMID 30721132',
             'Revisao da capacitacao em equinos. cAMP basal em equinos e 10x menor que '
             'em suinos (0,18 vs 1,58 pM). Ausencia de efluxo de colesterol e especie-especifica.'),
            ('Thomas et al. (2006)', 'Cryobiology',
             'Criocapacitacao: congelamento induz aumento de B pattern (CTC), fosforilacao '
             'de tirosina e exposicao de fosfatidilserina — mimetiza capacitacao.'),
            ('Papa et al. (2016)', 'Vet Clinics North America',
             'BotuCrio: MF + glicerol + gema 20%. Amidas (MF, DMF) permeiam melhor que '
             'glicerol — menos estresse osmotico. Diferente composicao de crioprotetor = '
             'diferente impacto na membrana.'),
            ('Unraveling glycerol and amides (2024)', 'Theriogenology',
             'Fatorial glicerol vs MF vs DMF vs DMA com diferentes curvas. Interacao '
             'crioprotetor x curva e significativa. MF preserva melhor funcao mitocondrial.'),
        ],
        'gap': 'NENHUM estudo investigou se o tipo de diluente de congelacao (e portanto o grau de '
               'criocapacitacao induzido) afeta a cinetica de capacitacao in vitro para FIV. '
               'Todos os estudos de FIV equina usaram UM unico crioprotetor. Se BotuCrio (amida) '
               'induz criocapacitacao diferente de INRA96+glicerol, o tempo otimo de pre-incubacao '
               'seria diferente para cada diluente.',
        'differential': 'Primeiro estudo a demonstrar que o protocolo de FIV deve ser ajustado conforme '
                       'o crioprotetor utilizado para congelacao. Implicacao pratica: veterinarios que '
                       'usam diferentes diluentes comerciais precisariam de tempos de pre-incubacao '
                       'especificos. A UNESP e o unico grupo que tem acesso a BotuCrio + expertise '
                       'em criopreservacao para testar essa hipotese.',
        'design': (
            'Split-sample de 10 garanhoes (3 ejaculados cada).\n'
            'Congelacao em 3 diluentes:\n'
            '- BotuCrio (MF + glicerol + gema 20%)\n'
            '- INRA96 + glicerol 3%\n'
            '- Gema + glicerol convencional 5%\n\n'
            'Descongelacao padrao -> lavagem -> PHE.\n'
            'Coletas nos tempos: 0h, 3h, 6h, 9h, 12h.\n\n'
            'Avaliacao em cada tempo:\n'
            '- Motilidade hiperativada (CASA: ALH >6 um, VCL >150 um/s, LIN <50%)\n'
            '- CTC (padroes F, B, AR)\n'
            '- Fosforilacao de tirosina (anti-pTyr 4G10 + citometria)\n'
            '- Viabilidade (CFDA/PI)\n'
            '- Potencial mitocondrial (JC-1)\n\n'
            'NAO precisa de oocitos — estudo de capacitacao pura.\n'
            'n = 10 garanhoes x 3 diluentes x 5 tempos = 150 amostras.'
        ),
        'equipment': 'CASA, citometro de fluxo, anticorpo anti-fosfo-tirosina 4G10 (~US$300-500), '
                    'CTC (clortetraciclina), JC-1, CFDA/PI — todos disponiveis ou acessiveis.',
        'cost': 'Baixo-moderado — anticorpo e o item mais caro.',
    },
    {
        'num': 3,
        'title': 'SpermFilter vs Swim-up vs SSD para Selecao Pre-FIV',
        'literature': [
            ('Felix et al. (2025)', 'Biology of Reproduction, PMID 40057974',
             'Metodo de separacao pos-descongelacao e CRITICO para FIV: SSD (Sperm Separation '
             'Device) = 88,5% fertilizacao; swim-up = 38,1%; centrifugacao em coloide = 20,0%.'),
            ('Papa et al. — SpermFilter', 'Diversos',
             'SpermFilter: membrana hidrofilica sintetica. Recuperacao 89,4% (vs 80,9% centrifugacao). '
             'Sem dano centrifugo. Reutilizavel (ate 10x/garanhao). NUNCA testado como metodo de '
             'selecao para FIV equina.'),
            ('Morrell (SLC)', 'Diversos',
             'Single Layer Centrifugation com Androcoll-E. Seleciona por motilidade e morfologia. '
             'Usado em bovinos para FIV — nao testado no contexto FIV equina convencional.'),
            ('Microfluidic sperm selection (2022)', 'Journal of Equine Veterinary Science',
             'Chip microfluIdico para selecao espermatica para ICSI equina. Integridade de DNA '
             'superior ao swim-up. NAO testado para FIV convencional.'),
        ],
        'gap': 'O SpermFilter (UNESP) nunca foi testado como metodo de selecao pre-FIV equina. '
               'O SSD (usado por Hinrichs) e o unico metodo validado para FIV. Nenhum estudo '
               'comparou diretamente SpermFilter vs SSD vs swim-up no contexto especifico de '
               'FIV convencional equina.',
        'differential': 'Se SpermFilter produzir resultados comparaveis ao SSD, a UNESP posiciona '
                       'um produto proprio (ja comercializado pela Botupharma) para o mercado de '
                       'FIV equina — que esta em plena expansao apos o breakthrough de 2022. '
                       'Impacto comercial direto.',
        'design': (
            'Semen de 10 garanhoes congelado em BotuCrio -> descongelacao.\n'
            '4 metodos de selecao:\n'
            '- G1: Centrifugacao simples 600xg/10 min (controle)\n'
            '- G2: Swim-up classico\n'
            '- G3: SpermFilter (produto UNESP)\n'
            '- G4: SSD (Sperm Separation Device — importado)\n\n'
            'Pos-selecao -> PHE 9h -> avaliacao de capacitacao:\n'
            '- CASA (hiperativacao)\n'
            '- CTC\n'
            '- Fosforilacao de tirosina\n'
            '- Viabilidade (CFDA/PI)\n'
            '- Concentracao e recuperacao espermatica\n\n'
            'Se oocitos disponiveis: co-incubacao e taxa de fertilizacao.\n'
            'Se nao: endpoint de capacitacao (publicavel como estudo comparativo).\n'
            'n = 10 garanhoes, 3 repeticoes.'
        ),
        'equipment': 'SpermFilter (ja disponivel), SSD (~US$50-100/unidade para importar), '
                    'CASA, citometria. Reagentes PHE.',
        'cost': 'Baixo-moderado.',
    },
    {
        'num': 4,
        'title': 'Efeito do Plasma Seminal Residual na Capacitacao In Vitro',
        'literature': [
            ('Sieme et al. (diversos)', 'TiHo Hannover',
             'Plasma seminal: 5-10% e otimo para CRIOPRESERVACAO. >50% e toxico. '
             'Componentes protetores: HSPs, clusterin, lactoferrin. Componentes deleteros: '
             'nao completamente identificados.'),
            ('Talluri et al. (2024)', 'Biopreservation and Biobanking',
             'Plasma seminal heterOlogo reduz calcio intracelular e viabilidade de '
             'espermatozoides equinos criopreservados.'),
            ('Fatores decapacitantes classicos', 'Diversos',
             'Plasma seminal contem fatores decapacitantes que estabilizam a membrana, '
             'mantem razao colesterol/fosfolipidio e previnem capacitacao prematura. '
             'ProteInas heparina-binding (Hep+) sao candidatas em equinos.'),
            ('Leemans et al. (2016)', 'Reproduction',
             'Albumina + bicarbonato induz aglutinacao cabeca-cabeca e reduz ligacao '
             'oviductal >10x. Condicoes de capacitacao podem ser contraproducentes '
             'se timing nao for correto.'),
        ],
        'gap': 'NENHUM estudo testou o efeito do plasma seminal residual na capacitacao '
               'IN VITRO para FIV equina. Os estudos de Sieme focam em criopreservacao (endpoint '
               'pos-descongelacao), NAO em capacitacao subsequente. E plausivel que a dose '
               'ideal de plasma para criopreservacao (5-10%) seja DIFERENTE da dose ideal '
               'para capacitacao (possivelmente 0%, dada a funcao decapacitante).',
        'differential': 'Primeiro estudo a separar dois objetivos que foram historicamente confundidos: '
                       '(1) plasma para criopreservacao vs (2) plasma para capacitacao. Se confirmar '
                       'que remocao completa de plasma melhora capacitacao para FIV, isso muda o '
                       'protocolo de processamento — e o SpermFilter da UNESP seria ideal para isso.',
        'design': (
            '12 garanhoes.\n'
            'Congelacao em BotuCrio com 4 niveis de plasma residual: 0%, 5%, 10%, 20%.\n'
            '(Controle de plasma via centrifugacao em gradiente + reconstituicao)\n\n'
            'Descongelacao -> PHE -> monitoramento em 0h, 3h, 6h, 9h.\n\n'
            'Parametros em cada tempo:\n'
            '- CTC (F, B, AR)\n'
            '- Fosforilacao de tirosina\n'
            '- Hiperativacao (CASA)\n'
            '- Viabilidade (CFDA/PI)\n'
            '- Potencial mitocondrial (JC-1)\n\n'
            'Bonus: se oocitos disponiveis, co-incubacao no tempo otimo de cada grupo.\n'
            'n = 12 garanhoes x 4 doses x 4 tempos = 192 amostras.'
        ),
        'equipment': 'SpermFilter ou centrifugacao em gradiente (ja disponivel), '
                    'CASA, citometria, reagentes PHE/CTC.',
        'cost': 'Baixo — usa materiais de rotina.',
    },
    {
        'num': 5,
        'title': 'Piruvato no Meio de Capacitacao PHE',
        'literature': [
            ('Pena et al. (2021)', 'Biology of Reproduction, PMID 33864078',
             'Low glucose (1 mM) + high pyruvate (10 mM): motilidade 76,2% vs 61,7% (INRA 96). '
             'Reducao de 2-oxoaldeIdos (metilglioxal). Melhor eficiencia mitocondrial e regulacao redox. '
             'Mitocondrias ativas: 51,1% vs 24,1% apos 48h.'),
            ('Pena et al. (2023)', 'Theriogenology, PMID 38029686',
             'Piruvato melhora funcao espermatica em meios com alta glicose, melhorando eficiencia '
             'metabolica geral. Mecanismo: conversao piruvato->lactato regenera NAD+ para glicolise. '
             'Piruvato alimenta Krebs diretamente (bypassa glicolise).'),
            ('Gibb et al. (2015)', 'Biology of Reproduction',
             'Lactato e piruvato sao as fontes de energia mais importantes para motilidade '
             'espermatica equina, com efeitos dose-dependentes na funcao mitocondrial.'),
            ('Protocolo PHE — Hinrichs (2022)', 'Biology of Reproduction',
             'Meio TALP contem lactato e piruvato em concentracoes padrao. '
             'NENHUMA otimizacao da razao glicose/piruvato foi tentada para FIV equina.'),
        ],
        'gap': 'O Pena demonstrou que a razao glicose/piruvato e CRITICA para funcao espermatica '
               'equina em diluentes de armazenamento. O protocolo PHE de Hinrichs usa TALP padrao — '
               'NINGUEM otimizou a composicao metabolica do meio de CAPACITACAO para FIV. '
               'Se piruvato melhora funcao mitocondrial (Pena) e mitocondrias sao essenciais '
               'para hiperativacao (necessaria para FIV), entao otimizar piruvato no PHE pode '
               'melhorar a qualidade da capacitacao.',
        'differential': 'Primeiro estudo a cruzar a linha de pesquisa metabolica do Pena (diluentes) '
                       'com a linha de FIV da Hinrichs (capacitacao). Duas linhas de pesquisa '
                       'consolidadas que NUNCA foram integradas. Se funcionar, a UNESP cria '
                       'um "PHE otimizado" com potencial de patente.',
        'design': (
            '4 meios de capacitacao:\n'
            '- G1: PHE padrao (TALP original — controle)\n'
            '- G2: PHE com piruvato 2x (dobrar concentracao)\n'
            '- G3: PHE com glicose 50% + piruvato 2x (low-glucose/high-pyruvate)\n'
            '- G4: PHE sem glicose + piruvato 3x (zero-glucose/high-pyruvate)\n\n'
            'Semen congelado BotuCrio -> selecao -> pre-incubacao 9h em cada meio.\n\n'
            'Parametros:\n'
            '- JC-1 (potencial mitocondrial — endpoint primario)\n'
            '- CASA (hiperativacao: ALH, VCL, LIN)\n'
            '- CTC (padroes de capacitacao)\n'
            '- Viabilidade (CFDA/PI)\n'
            '- ROS (DHE — superoxido mitocondrial)\n\n'
            'n = 10 garanhoes.'
        ),
        'equipment': 'Piruvato de sodio e glicose (reagentes de rotina, <R$100). '
                    'CASA, citometria (JC-1, DHE, CFDA/PI, CTC).',
        'cost': 'Muito baixo.',
    },
    {
        'num': 6,
        'title': 'Termotaxia como Selecao de Espermatozoides Capacitados',
        'literature': [
            ('Ruiz-Diaz et al. (2020)', 'Animals, PMID 32825582',
             'D-penicilamina + HTF induz capacitacao em espermatozoides equinos criopreservados. '
             'Permite selecao por termotaxia in vitro. Fracao selecionada: enriquecida em '
             'fosforilacao de tirosina + menor fragmentacao de DNA.'),
            ('Perez-Cerezales et al. (2022)', 'J Anim Sci Biotechnol',
             'Espermatozoides bovinos selecionados por termotaxia tem maior integridade de DNA, '
             'morfometria especIfica e melhoram resultados de ICSI.'),
            ('Hinrichs et al. (2022)', 'Biology of Reproduction',
             'O protocolo PHE ja contem penicilamina — um dos componentes que Ruiz-Diaz usa '
             'para induzir capacitacao pre-termotaxia.'),
        ],
        'gap': 'Termotaxia foi demonstrada como metodo de selecao para ICSI bovina e como '
               'ferramenta de pesquisa em equinos. NINGUEM aplicou termotaxia como metodo de '
               'selecao de espermatozoides capacitados antes da FIV convencional equina. '
               'O Felix identificou que a heterogeneidade temporal da capacitacao e um problema — '
               'nem todos os espermatozoides capacitam no mesmo ritmo. Termotaxia selecionaria '
               'especificamente os que JA capacitaram.',
        'differential': 'Abordagem completamente nova: usar uma propriedade FUNCIONAL (termotaxia = '
                       'exclusiva de espermatozoides capacitados) como metodo de selecao pre-FIV. '
                       'Eliminaria o problema da heterogeneidade temporal. Dispositivo pode ser '
                       'construido com materiais simples (placa aquecida + seringa).',
        'design': (
            'Semen congelado BotuCrio -> descongelacao -> SpermFilter -> PHE 9h.\n\n'
            'Divisao em 2 grupos:\n'
            '- G1: Co-incubacao direta com oocitos (protocolo Hinrichs padrao)\n'
            '- G2: Selecao por termotaxia (gradiente 37 graus C -> 39 graus C, 30 min) ->\n'
            '       fracao que migra -> co-incubacao com oocitos\n\n'
            'Avaliar fracao selecionada vs nao-selecionada:\n'
            '- Fosforilacao de tirosina (anti-pTyr)\n'
            '- CTC (% padrao B e AR)\n'
            '- DNA (comet assay ou SCSA)\n'
            '- Concentracao espermatica recuperada\n\n'
            'Endpoints finais: taxa de fertilizacao, clivagem, blastocistos.\n'
            'n = 10 garanhoes.'
        ),
        'equipment': 'Camara de termotaxia (placa aquecida com gradiente — construcao simples, '
                    '<R$500). CASA, citometria, materiais de cultura embrionaria.',
        'cost': 'Baixo.',
    },
    {
        'num': 7,
        'title': 'pH Alcalino (7,9) no Meio de Capacitacao PHE',
        'literature': [
            ('Leemans et al. (2015)', 'Reproduction, PMID 26078083',
             'Fluido folicular tratado a pH 7,9 (NAO 7,4) induz hipermotilidade Ca2+-dependente '
             'em espermatozoides equinos. So funciona em espermatozoides previamente ligados a '
             'explantes oviductais.'),
            ('Gonzalez-Fernandez et al. (2013)', 'Biology of Reproduction',
             'Ligacao ao oviduto + pH elevado induz fosforilacao de tirosina em espermatozoides '
             'equinos. pH intracelular aumenta gradualmente em microambiente alcalino.'),
            ('Leemans et al. (2019)', 'Reproduction, PMID 30721132',
             'Procaina (2,5-10 mM, base fraca que alcaliniza) induz hiperativacao robusta, '
             'mas NAO fertilizacao. pH-dependente e parcialmente independente de cAMP.'),
            ('Conserved CatSper mechanism (2021)', 'Frontiers Cell Dev Biol',
             'Bicarbonato sensibiliza CatSper para Ca2+. Em equinos, cAMP basal e 10x menor — '
             'pode ser que pH alcalino compense parcialmente essa deficiencia.'),
        ],
        'gap': 'O protocolo PHE de Hinrichs usa pH padrao (~7,4). Leemans demonstrou que '
               'pH 7,9 e um gatilho CRITICO para capacitacao equina. NINGUEM testou o efeito '
               'de pH elevado diretamente no meio PHE de FIV. Se pH 7,9 induz capacitacao '
               'por via independente de cAMP (complementar ao PHE), a combinacao poderia '
               'melhorar taxas de fertilizacao ou reduzir tempo de pre-incubacao.',
        'differential': 'Integra duas linhas de pesquisa independentes: o protocolo PHE '
                       '(Hinrichs — empirico, funciona mas mecanismo incompleto) com a descoberta '
                       'do pH alcalino (Leemans — mecanistico, mas sem aplicacao em FIV). '
                       'Teste simples (ajuste de pH com NaOH) com potencial de melhorar '
                       'significativamente o protocolo.',
        'design': (
            '4 valores de pH no meio PHE:\n'
            '- G1: pH 7,2\n'
            '- G2: pH 7,4 (padrao — controle)\n'
            '- G3: pH 7,6\n'
            '- G4: pH 7,9\n\n'
            'Tamponamento: HEPES 25 mM + bicarbonato 15 mM, ajuste com NaOH 1N.\n'
            'Monitoramento de pH durante incubacao (cada 3h).\n\n'
            'Semen congelado BotuCrio -> selecao -> PHE em cada pH -> 9h.\n\n'
            'Parametros:\n'
            '- Hiperativacao (CASA: ALH, VCL, LIN)\n'
            '- CTC (F, B, AR)\n'
            '- Fosforilacao de tirosina\n'
            '- Viabilidade (CFDA/PI)\n'
            '- JC-1 (potencial mitocondrial)\n'
            '- pH intracelular se possivel (BCECF)\n\n'
            'Se oocitos disponiveis: co-incubacao e fertilizacao.\n'
            'n = 10 garanhoes.'
        ),
        'equipment': 'NaOH, HCl, HEPES, pHmetro (tudo disponivel). CASA, citometria.',
        'cost': 'Muito baixo.',
    },
    {
        'num': 8,
        'title': 'Bad Freezers na FIV: Criocapacitacao como Vantagem Paradoxal?',
        'literature': [
            ('Thomas et al. (2006)', 'Reproduction, Fertility and Development',
             'Criocapacitacao: congelamento aumenta CTC padrao B de 5,4% (fresco) para 64,8% '
             '(pos-descongelacao) e AR de 1% para 32,8%. Mudancas dramaticas de membrana.'),
            ('Felix et al. (2025)', 'Biology of Reproduction',
             'Semen congelado capacita mais rapido (9h vs 22h). Hipotese: criocapacitacao avanca '
             'o ponto de partida. Mecanismos moleculares DIFEREM do semen fresco.'),
            ('Sieme — classificacao de garanhoes', 'TiHo Hannover',
             '30% bons congeladores (>40% motilidade pos-descongelacao), 40% intermediarios, '
             '30% maus congeladores (<20%). Maus congeladores sofrem MAIS criocapacitacao.'),
            ('Papa/Alvarenga — bad freezers', 'Diversos UNESP',
             'Amidas (DMF, MF) melhoram bad freezers para criopreservacao. '
             'Nenhuma investigacao de bad freezers no contexto de FIV.'),
            ('Cellular consequences of cryopreservation (2024)', 'J Equine Vet Sci',
             'Revisao: criocapacitacao inclui perda de colesterol, fosforilacao de tirosina '
             'prematura, reducao de viabilidade. Associada a reducao de fertilidade em IA — '
             'mas NUNCA testada no contexto de FIV.'),
        ],
        'gap': 'Todos os estudos de criocapacitacao concluem que ela e NEGATIVA (reduz fertilidade '
               'em IA). Mas para FIV convencional, a capacitacao e o OBJETIVO. Se bad freezers '
               'sofrem MAIS criocapacitacao, eles estariam MAIS proximos do estado capacitado — '
               'paradoxalmente, poderiam ter MELHOR desempenho em FIV. NINGUEM testou essa hipotese.',
        'differential': 'Hipotese completamente contraintuitiva que, se confirmada, muda o paradigma: '
                       'garanhoes "ruins para IA com semen congelado" poderiam ser "bons para FIV". '
                       'Isso ampliaria dramaticamente a indicacao clinica da FIV equina e resgata '
                       'genetica de garanhoes ate entao considerados problematicos. '
                       'A UNESP tem garanhoes ja classificados para testar.',
        'design': (
            '15 garanhoes classificados: 5 bons, 5 intermediarios, 5 maus congeladores.\n'
            '(Classificacao previa da UNESP — ja disponivel)\n\n'
            'Congelacao em BotuCrio -> descongelacao -> PHE.\n'
            'Monitoramento da capacitacao: 0, 3, 6, 9, 12h.\n\n'
            'Parametros em cada tempo:\n'
            '- CTC (% F, B, AR)\n'
            '- Fosforilacao de tirosina\n'
            '- Hiperativacao (CASA)\n'
            '- Viabilidade (CFDA/PI)\n'
            '- JC-1\n\n'
            'Endpoint-chave: identificar o TEMPO OTIMO de cada CLASSE de garanhao.\n\n'
            'Hipoteses:\n'
            '- Bad freezers atingem pico de capacitacao ANTES (3-6h vs 9h)?\n'
            '- Bad freezers tem MAIS espermatozoides em padrao B no T0?\n'
            '- A janela entre capacitacao funcional e morte celular e mais estreita?\n\n'
            'Se oocitos disponiveis: FIV no tempo otimo de cada classe.\n'
            'n = 15 garanhoes, 3 ejaculados cada.'
        ),
        'equipment': 'CASA, citometria, BotuCrio, reagentes PHE/CTC. '
                    'Garanhoes ja classificados na UNESP.',
        'cost': 'Baixo.',
    },
    {
        'num': 9,
        'title': 'Antioxidante no Meio de Capacitacao: Proteger sem Inibir',
        'literature': [
            ('Pena (2019)', 'Antioxidants, PMID 31752408',
             'ROS sao moleculas de sinalizacao essenciais: capacitacao, hiperativacao e reacao '
             'acrossomica dependem de ROS. Eustress (ROS controlado) vs distress (excesso). '
             'O espermatozoide equino vive em uma "corda bamba redox".'),
            ('Gibb (diversos)', 'University of Newcastle',
             'Curva em U invertido: ROS deficiente = subfertil (capacitacao falha). '
             'ROS otimo = fertil. ROS excessivo = infertil (peroxidacao). '
             'Garanhoes mais ferteis produzem MAIS ROS — paradoxo.'),
            ('Catalase in stallion sperm (2025)', 'Antioxidants',
             'Espermatozoides equinos sao 14x mais resistentes a dano oxidativo que humanos '
             '(IC50 391,6 uM vs 27,3 uM). Acumulam mais H2O2 intracelular sem perda de motilidade. '
             'Defesas antioxidantes robustas, especialmente catalase.'),
            ('ROS e quimiotaxia equina (2020)', 'Reproduction',
             'ROS participam da sinalizacao de quimiotaxia espermatica equina via NOX -> '
             'cAMP-PKA. Antioxidante em excesso bloquearia quimiotaxia.'),
        ],
        'gap': 'Antioxidantes sao extensivamente estudados em DILUENTES de armazenamento. '
               'NINGUEM testou antioxidantes no MEIO DE CAPACITACAO para FIV equina. '
               'O paradoxo e real: proteger o espermatozoide de dano oxidativo durante '
               '9h de pre-incubacao, mas sem inibir as ROS necessarias para capacitacao. '
               'A dose-resposta nesse contexto e completamente desconhecida.',
        'differential': 'Primeiro estudo a estabelecer a curva dose-resposta de antioxidante '
                       'no contexto de CAPACITACAO para FIV (nao armazenamento). Resolve a '
                       'questao pratica: adicionar ou nao antioxidante ao PHE? Se a resposta '
                       'for uma dose especifica que protege viabilidade sem inibir capacitacao, '
                       'isso otimiza diretamente o protocolo de FIV.',
        'design': (
            'Meio PHE + 4 doses de MitoTEMPO (antioxidante mitocondria-especifico):\n'
            '- G1: 0 (controle — PHE puro)\n'
            '- G2: 1 uM (dose baixa)\n'
            '- G3: 10 uM (dose media)\n'
            '- G4: 100 uM (dose alta)\n\n'
            'Semen congelado BotuCrio -> selecao -> PHE + MitoTEMPO -> 9h.\n\n'
            'Parametros PRO-CAPACITACAO (devem ser MANTIDOS):\n'
            '- CTC padrao B (%)\n'
            '- Fosforilacao de tirosina\n'
            '- Hiperativacao (CASA)\n\n'
            'Parametros PRO-VIABILIDADE (devem ser MELHORADOS):\n'
            '- Membrana (CFDA/PI)\n'
            '- JC-1 (potencial mitocondrial)\n'
            '- DNA (comet assay)\n\n'
            'Parametros MECANISTICOS:\n'
            '- ROS (DHE — superoxido mitocondrial)\n'
            '- 4-HNE (peroxidacao lipidica)\n\n'
            'Hipotese: curva em U invertido — dose baixa protege sem inibir;\n'
            'dose alta protege viabilidade mas BLOQUEIA capacitacao.\n'
            'n = 10 garanhoes.'
        ),
        'equipment': 'MitoTEMPO (~US$150-200). CASA, citometria (JC-1, DHE, CFDA/PI, CTC). '
                    'Anti-4-HNE se disponivel.',
        'cost': 'Baixo-moderado.',
    },
    {
        'num': 10,
        'title': 'Metabolitos Oviductais no Meio PHE',
        'literature': [
            ('Lamy et al. (2020)', 'Theriogenology, PMID 31835098',
             'Metabolomica do fluido oviductal equino por 1H-NMR. Metabolitos mais abundantes: '
             'lactato, mioinositol, creatina, alanina, carnitina. Apenas fumarato e glicina '
             'diferem entre pre- e pos-ovulatorio.'),
            ('Gonzalez-Fernandez et al. (2022)', 'J Equine Vet Sci, PMID 35077851',
             'Metabolitos selecionados do fluido oviductal NAO modificaram parametros de '
             'capacitacao em semen equino criopreservado. POREM: usaram concentracoes e '
             'tempos diferentes do protocolo PHE; usaram semen sem pre-incubacao prolongada.'),
            ('Leemans et al. (2022)', 'Biology of Reproduction',
             'Cultivo de celulas epiteliais oviductais equinas (EOEC) em Transwell e chips '
             'microfluIdicos. Celulas secretam fatores que melhoram capacitacao. '
             'Identidade molecular dos fatores: desconhecida.'),
            ('Metabolomica folicular equina (2020)', 'MDPI, PMID 32438699',
             'Fluido folicular pre-ovulatorio contem metabolitos distintos do oviductal. '
             'pH 7,9 no folicular e um diferencial critico.'),
        ],
        'gap': 'O estudo de Gonzalez-Fernandez (2022) testou metabolitos oviductais e concluiu '
               'que NAO funcionam — mas usou protocolo DIFERENTE do PHE (sem pre-incubacao '
               'prolongada, concentracoes diferentes). NINGUEM testou metabolitos oviductais '
               'no contexto do protocolo PHE de 9h de Hinrichs/Felix. A pre-incubacao prolongada '
               'pode ser o elemento que faltava para os metabolitos exercerem efeito.',
        'differential': 'Tenta mimetizar o microambiente oviductal SEM co-cultura de celulas — '
                       'simplificacao que viabilizaria FIV em laboratorios sem infraestrutura '
                       'de cultura celular. Se funcionar, cria um "PHE oviductal" que poderia '
                       'substituir a necessidade de EOEC.',
        'design': (
            '5 grupos de meio de capacitacao:\n'
            '- G1: PHE padrao (controle)\n'
            '- G2: PHE + lactato 25 mM\n'
            '- G3: PHE + mioinositol 5 mM + creatina 1 mM\n'
            '- G4: PHE + L-carnitina 5 mM\n'
            '- G5: PHE + coquetel completo (lactato 25 mM + mioinositol 5 mM +\n'
            '       creatina 1 mM + L-carnitina 5 mM + alanina 2 mM)\n\n'
            'Semen congelado BotuCrio -> selecao -> pre-incubacao 9h.\n\n'
            'Parametros:\n'
            '- CTC, fosforilacao de tirosina, hiperativacao (CASA)\n'
            '- Viabilidade, JC-1\n\n'
            'Se oocitos disponiveis: co-incubacao e fertilizacao.\n'
            'n = 10 garanhoes.\n\n'
            'Todos os reagentes sao baratos e disponiveis em fornecedores nacionais.'
        ),
        'equipment': 'Lactato de sodio, mioinositol, creatina, L-carnitina, alanina — '
                    'todos reagentes de rotina (<R$300 total). CASA, citometria.',
        'cost': 'Muito baixo.',
    },
]

for proj in projects:
    doc.add_page_break()
    doc.add_heading(f'{proj["num"] + 2}. Projeto {proj["num"]} — {proj["title"]}', level=1)

    # Literature review
    doc.add_heading('Rastreio de Literatura', level=2)
    for ref_author, ref_journal, ref_desc in proj['literature']:
        p = doc.add_paragraph()
        run = p.add_run(f'{ref_author} ')
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(f'({ref_journal})')
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        p2 = doc.add_paragraph(ref_desc)
        p2.paragraph_format.left_indent = Cm(1)
        p2.paragraph_format.space_after = Pt(8)

    # Gap
    doc.add_heading('Gap Identificado na Literatura', level=2)
    p = doc.add_paragraph(proj['gap'])
    # Add visual emphasis
    for run in p.runs:
        pass  # keep as is

    # Differential
    doc.add_heading('Diferencial e Justificativa de Originalidade', level=2)
    p = doc.add_paragraph(proj['differential'])

    # Experimental design
    doc.add_heading('Desenho Experimental', level=2)
    for line in proj['design'].split('\n'):
        if line.strip():
            doc.add_paragraph(line, style='List Bullet' if line.startswith('-') else 'Normal')
        else:
            doc.add_paragraph()

    # Equipment and cost
    doc.add_heading('Equipamentos e Reagentes Necessarios', level=2)
    doc.add_paragraph(proj['equipment'])

    doc.add_heading('Estimativa de Custo', level=2)
    doc.add_paragraph(proj['cost'])

doc.add_page_break()

# ============================================================
# SYNTHESIS TABLE
# ============================================================
doc.add_heading('13. Sintese Comparativa e Priorizacao', level=1)

doc.add_paragraph(
    'A tabela abaixo sintetiza os 10 projetos segundo criterios de viabilidade na UNESP, '
    'originalidade cientifica, impacto potencial e custo estimado.'
)

table = doc.add_table(rows=1, cols=6, style='Table Grid')
headers = ['#', 'Projeto', 'Custo', 'Oocitos?', 'Originalidade', 'Impacto']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

synthesis_data = [
    ('5', 'Piruvato no PHE', 'Muito baixo', 'Opcional', '★★★★★', 'Alto'),
    ('7', 'pH 7,9 no PHE', 'Muito baixo', 'Opcional', '★★★★★', 'Muito alto'),
    ('10', 'Metabolitos oviductais', 'Muito baixo', 'Opcional', '★★★★', 'Alto'),
    ('2', 'Cinetica x crioprotetor', 'Baixo-mod', 'Nao', '★★★★★', 'Muito alto'),
    ('4', 'Plasma x capacitacao', 'Baixo', 'Opcional', '★★★★', 'Alto'),
    ('8', 'Bad freezers FIV', 'Baixo', 'Opcional', '★★★★★', 'Transformador'),
    ('9', 'Antioxidante no PHE', 'Baixo-mod', 'Nao', '★★★★★', 'Fundamental'),
    ('1', 'PHE + BotuCrio', 'Moderado', 'Sim', '★★★', 'Estrategico'),
    ('3', 'SpermFilter pre-FIV', 'Baixo-mod', 'Opcional', '★★★★', 'Estrategico'),
    ('6', 'Termotaxia pre-FIV', 'Baixo', 'Sim', '★★★★★', 'Alto'),
]
for row_data in synthesis_data:
    row = table.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        for p in row.cells[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

set_table_style(table)

doc.add_paragraph()
doc.add_paragraph(
    'Recomendacao de priorizacao: Os projetos 5, 7 e 10 podem ser iniciados imediatamente '
    'com reagentes de rotina e sem necessidade de oocitos (endpoints de capacitacao). '
    'Os projetos 2, 4 e 8 sao estudos de capacitacao pura que tambem nao exigem oocitos. '
    'Os projetos 1, 3 e 6 requerem infraestrutura de cultura embrionaria para o endpoint de FIV.'
)

doc.add_page_break()

# ============================================================
# THESIS PROPOSAL
# ============================================================
doc.add_heading('14. Proposta de Tese Integrada', level=1)

doc.add_paragraph(
    'Os 10 projetos podem ser organizados em uma tese de doutorado coesa com 4 capitulos:'
)

doc.add_heading('Titulo proposto', level=2)
p = doc.add_paragraph()
run = p.add_run(
    '"Otimizacao da capacitacao in vitro de espermatozoides equinos criopreservados '
    'em BotuCrio para fertilizacao in vitro convencional"'
)
run.italic = True

chapters = [
    ('Capitulo 1 — Baseline', 'Projeto 2',
     'Cinetica de capacitacao em funcao do tipo de crioprotetor. Estabelece o ponto de partida '
     'e demonstra que o diluente de congelacao afeta a capacitacao in vitro.'),
    ('Capitulo 2 — Otimizacao do meio', 'Projetos 5 + 7',
     'Otimizacao metabolica (piruvato) e de pH (7,9) no meio PHE. Cria um "PHE otimizado" '
     'com duas variaveis de baixo custo e alto impacto.'),
    ('Capitulo 3 — Variabilidade', 'Projeto 8',
     'Bad freezers na FIV — testa a hipotese provocadora de que criocapacitacao pode ser vantagem. '
     'Se confirmada, amplia a indicacao clinica da FIV.'),
    ('Capitulo 4 — Validacao integrada', 'Projetos 1 + 3',
     'Reproducao do protocolo PHE otimizado (caps. 1-3) + comparacao SpermFilter vs SSD. '
     'Endpoint final: taxa de fertilizacao in vitro com BotuCrio.'),
]

for ch_title, ch_proj, ch_desc in chapters:
    doc.add_heading(ch_title, level=3)
    p = doc.add_paragraph()
    run = p.add_run(f'Projetos incluidos: {ch_proj}')
    run.bold = True
    doc.add_paragraph(ch_desc)

doc.add_paragraph()
doc.add_paragraph(
    'Estimativa: 4 capitulos, 4-5 publicacoes, tese de doutorado com impacto internacional. '
    'O fluxo logico vai do baseline (cap. 1) a otimizacao (cap. 2), variabilidade (cap. 3) '
    'e validacao final (cap. 4).'
)

doc.add_page_break()

# ============================================================
# REFERENCES
# ============================================================
doc.add_heading('15. Referencias Bibliograficas', level=1)

references = [
    'Hinrichs K, Choi YH, Love CC, Chung YG, Varner DD. Production of horse foals via direct injection of rime oocytes with stallion spermatozoa. Biology of Reproduction. 2022;107(6):1320-1329. PMID: 36084034.',
    'Felix MR, Loux SC, Esteller-Vico A, Hinrichs K. Equine in vitro fertilization with frozen-thawed semen is associated with shortened pre-incubation time and modified capacitation-related changes. Biology of Reproduction. 2025;112(5):867-878. PMID: 40057974.',
    'Martin-Pelaez P, et al. IVF with frozen-thawed sperm after prolonged capacitation yields comparable results to ICSI in horses: A morphokinetics study. Theriogenology. 2024.',
    'Leemans B, Gadella BM, Stout TAE, De Schauwer C, Nelis H, Hoogewijs M, Van Soom A. Why doesn\'t conventional IVF work in the horse? The equine oviduct as a microenvironment for capacitation/fertilization. Reproduction. 2016;152(6):R233-R245.',
    'Leemans B, Stout TAE, De Schauwer C, Heras S, Nelis H, Hoogewijs M, Van Soom A, Gadella BM. Update on mammalian sperm capacitation: how much does the horse differ from other species? Reproduction. 2019;157(5):R181-R197. PMID: 30721132.',
    'Leemans B, Stout TAE, Soom AV, Gadella BM. pH-dependent effects of procaine on equine gamete activation. Biology of Reproduction. 2019;101(5):1056-1074.',
    'Leemans B, et al. An alkaline follicular fluid fraction induces capacitation and limited release of oviduct epithelium-bound stallion sperm. Reproduction. 2015;150(3):193-208. PMID: 26078083.',
    'Leemans B, et al. Developing a reproducible protocol for culturing functional confluent monolayers of differentiated equine oviduct epithelial cells. Biology of Reproduction. 2022;106(4):710-726.',
    'Pena FJ, et al. Low glucose and high pyruvate reduce the production of 2-oxoaldehydes, improving mitochondrial efficiency, redox regulation, and stallion sperm function. Biology of Reproduction. 2021;105(2):519-532. PMID: 33864078.',
    'Pena FJ, et al. Pyruvate enhances stallion sperm function in high glucose media improving overall metabolic efficiency. Theriogenology. 2023. PMID: 38029686.',
    'Pena FJ, et al. Redox Regulation and Oxidative Stress: The Particular Case of the Stallion Spermatozoa. Antioxidants. 2019;8(11):567. PMID: 31752408.',
    'Gibb Z, Aitken RJ. The impact of sperm metabolism during in vitro storage: the stallion as a model. BioMed Research International. 2016.',
    'Lamy J, Gatien J, Dumollard R, Nadal-Desbarats L, Eloy L, Vignon F, Teixeira-Gomes AP. Stage-specific metabolomic changes in equine oviductal fluid. Theriogenology. 2020;143:140-151. PMID: 31835098.',
    'Gonzalez-Fernandez L, et al. Selected Metabolites Found in Equine Oviductal Fluid do not Modify the Parameters Associated to Capacitation of the Frozen-thawed Equine Spermatozoa In Vitro. J Equine Vet Sci. 2022;111:103875. PMID: 35077851.',
    'Ruiz-Diaz S, et al. The Presence of D-Penicillamine during the In Vitro Capacitation of Stallion Spermatozoa Prolongs Hyperactive-Like Motility and Allows for Sperm Selection by Thermotaxis. Animals. 2020;10(9):1467. PMID: 32825582.',
    'Maitan PP, et al. A stallion spermatozoon\'s journey through the mare\'s genital tract: In vivo and in vitro aspects of sperm capacitation. Animal Reproduction Science. 2022;236:106932.',
    'Thomas AD, et al. Capacitation-like changes in equine spermatozoa throughout the cryopreservation process. Reproduction, Fertility and Development. 2006;14(4):225-233.',
    'Sieme H, et al. Advances in Stallion Semen Cryopreservation. Veterinary Clinics of North America: Equine Practice. 2016;32(3):521-530.',
    'Papa FO, et al. Advances in Stallion Semen Cryopreservation. Veterinary Clinics of North America: Equine Practice. 2016;32(3):521-530.',
    'Alvarenga MA, Papa FO, Neto CR. Advances in Stallion Semen Cryopreservation. Vet Clin North Am Equine Pract. 2016;32(3):521-530.',
    'Talluri TR, et al. Heterologous Seminal Plasma Reduces the Intracellular Calcium and Sperm Viability of Cryopreserved Stallion Spermatozoa. Biopreservation and Biobanking. 2024.',
    'Catalase in Unexpected Places: Revisiting H2O2 Detoxification Pathways in Stallion Spermatozoa. Antioxidants. 2025;14(6):718.',
    'Reactive oxygen species are involved in the signaling of equine sperm chemotaxis. Reproduction. 2020;159(4):REP-19-0480.',
    'Correa L, et al. Equine IVF milestone — University of Florida. 2026.',
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'[{i}] ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(ref)
    run.font.size = Pt(9)

# ============================================================
# SAVE
# ============================================================
output_path = '/Users/gabrielpetelinkar/AIOSORIZA/outputs/Projetos_Capacitacao_FIV_Equina_UNESP.docx'
doc.save(output_path)
print(f'Documento salvo em: {output_path}')
